#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
人事管理系统后端API
功能：用户认证、工资数据管理、签名文件管理、请假管理
"""

from flask import Flask, request, jsonify, send_file, render_template, redirect
from flask_cors import CORS
import sqlite3
import os
import sys
import json
import base64
import threading
import uuid
from datetime import datetime
from werkzeug.security import generate_password_hash
from werkzeug.utils import secure_filename

# 导出任务管理（数据库存储，支持多worker共享）
def _get_export_task(task_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM export_tasks WHERE task_id = ?", (task_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return {
        'task_id': row['task_id'],
        'status': row['status'],
        'progress': row['progress'],
        'total': row['total'],
        'message': row['message'],
        'file_path': row['file_path'],
        'filename': row['filename'],
        'error': row['error'],
        'created_at': row['created_at']
    }

def _set_export_task(task_id, **kwargs):
    conn = get_db_connection()
    cursor = conn.cursor()
    fields = []
    values = []
    for k, v in kwargs.items():
        if k in ('status', 'progress', 'total', 'message', 'file_path', 'filename', 'error'):
            fields.append(f"{k} = ?")
            values.append(v)
    if fields:
        values.append(task_id)
        cursor.execute(f"UPDATE export_tasks SET {', '.join(fields)} WHERE task_id = ?", values)
        conn.commit()
    conn.close()

def _create_export_task(task_id, data):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO export_tasks (task_id, status, progress, total, message, file_path, filename, error, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (task_id, data.get('status', 'pending'), data.get('progress', 0), data.get('total', 0),
         data.get('message', ''), data.get('file_path'), data.get('filename'), data.get('error'),
         data.get('created_at', datetime.now().isoformat()))
    )
    conn.commit()
    conn.close()

# 确保标准输出使用UTF-8编码
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

from auth import AuthManager, login_required, admin_required
from salary_manager import SalaryManager
from signature_manager import SignatureManager
from data_manager import DataManager
from device_manager import DeviceManager

app = Flask(__name__)
app.config['TEMPLATES_AUTO_RELOAD'] = True  # 禁用模板缓存
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # 禁用静态文件缓存
app.jinja_env.auto_reload = True  # 强制Jinja2重新加载模板
CORS(app)

DATABASE_PATH = os.path.join(os.path.dirname(__file__), '..', 'database', 'salary_system.db')
SIGNATURE_DIR = os.path.join(os.path.dirname(__file__), 'signatures')


def get_db_connection():
    """获取数据库连接"""
    conn = sqlite3.connect(DATABASE_PATH, timeout=60)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=60000")
    conn.execute("PRAGMA wal_checkpoint(PASSIVE)")
    return conn


def ensure_signature_dir():
    """确保签名目录存在"""
    if not os.path.exists(SIGNATURE_DIR):
        os.makedirs(SIGNATURE_DIR)


# ========================================
# Web界面路由
# ========================================

@app.route('/')
def index():
    """首页 - 重定向到登录页"""
    return redirect('/login')


@app.route('/login')
def login_page():
    """登录页面"""
    return render_template('login.html')


@app.route('/admin')
def admin_page():
    """管理员页面 - 工资条管理"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, 'templates', 'admin.html')
    
    print(f"模板文件路径: {template_path}")
    print(f"文件存在: {os.path.exists(template_path)}")
    
    if os.path.exists(template_path):
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        print(f"文件大小: {len(template_content)}")
        print(f"包含月份输入框: {'工资表所属月份' in template_content}")
        
        return template_content
    else:
        return "模板文件不存在", 500


@app.route('/menu')
def menu_page():
    """功能菜单页面"""
    return render_template('menu.html')


@app.route('/user-manage')
def user_manage_page():
    """人员管理页面"""
    return render_template('user_manage.html')


@app.route('/device-manage')
def device_manage_page():
    """设备管理页面"""
    return render_template('device_manage.html')


@app.route('/cleanup')
def cleanup_page():
    """数据清理页面"""
    return render_template('cleanup.html')


@app.route('/leave-manage')
def leave_manage_page():
    """请假管理页面"""
    return render_template('leave_manage.html')


@app.route('/customer-archive')
def customer_archive_page():
    """客户档案页面"""
    return render_template('customer_archive.html')


@app.route('/app-update')
def app_update_page():
    """APP更新管理页面"""
    return render_template('app_update.html')


# ========================================
# API路由
# ========================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查端点"""
    try:
        conn = get_db_connection()
        conn.execute("PRAGMA wal_checkpoint(TRUNCATE)")
        conn.close()
    except:
        pass
    
    return jsonify({
        'status': 'ok',
        'message': '人事管理系统后端服务运行正常',
        'timestamp': datetime.now().isoformat()
    })


# APP 版本信息（可在此处修改最新版本号）
APP_VERSION_CODE = 2
APP_VERSION_NAME = "1.0.2"
APP_DOWNLOAD_URL = "http://yaohu.dynv6.net:32996/app/download"
APP_UPDATE_MESSAGE = "发现新版本，请及时更新以获得最佳体验"

@app.route('/api/app/version', methods=['GET'])
def get_app_version():
    """获取APP最新版本信息"""
    return jsonify({
        'success': True,
        'version_code': APP_VERSION_CODE,
        'version_name': APP_VERSION_NAME,
        'download_url': APP_DOWNLOAD_URL,
        'update_message': APP_UPDATE_MESSAGE
    })


@app.route('/app/download', methods=['GET'])
def download_app():
    """下载APK文件"""
    apk_path = os.path.join(os.path.dirname(__file__), '..', 'app-release.apk')
    if not os.path.exists(apk_path):
        return jsonify({'success': False, 'error': 'APK文件不存在'}), 404
    return send_file(apk_path, as_attachment=True, download_name='salary-system.apk', mimetype='application/vnd.android.package-archive')


@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    
    if not data:
        return jsonify({
            'success': False,
            'error': '请求体为空'
        }), 400
    
    name = data.get('name')
    password = data.get('password')
    device_id = data.get('device_id', '')
    device_info = data.get('device_info', '')
    
    if not name or not password:
        return jsonify({
            'success': False,
            'error': '姓名和密码不能为空'
        }), 400
    
    auth_manager = AuthManager(DATABASE_PATH)
    result = auth_manager.login(name, password)
    
    if not result['success']:
        return jsonify(result), 401
    
    if result['role'] != 'admin' and device_id:
        device_manager = DeviceManager(DATABASE_PATH)
        
        verified, message, matched_username = device_manager.verify_device(name, device_id)
        
        if not verified:
            device_manager.record_login_attempt(
                username=name,
                device_id=device_id,
                device_info=device_info,
                status='failed',
                matched_username=matched_username,
                ip_address=request.remote_addr
            )
            
            return jsonify({
                'success': False,
                'error': '设备ID与账号不符，拒绝登录',
                'need_bind': False
            }), 403
        
        if verified and "首次登录" in message:
            bind_success, bind_message = device_manager.bind_device(name, device_id, device_info)
            if bind_success:
                result['device_bound'] = True
            else:
                return jsonify({
                    'success': False,
                    'error': bind_message
                }), 400
        
        device_manager.record_login_attempt(
            username=name,
            device_id=device_id,
            device_info=device_info,
            status='success',
            ip_address=request.remote_addr
        )
    
    return jsonify(result)


@app.route('/api/logout', methods=['POST'])
@login_required
def logout():
    """
    用户登出
    需要在Header中提供Authorization: token
    """
    auth_manager = AuthManager(DATABASE_PATH)
    success = auth_manager.logout(request.user_name)
    
    if success:
        return jsonify({
            'success': True,
            'message': '登出成功'
        })
    else:
        return jsonify({
            'success': False,
            'error': '登出失败'
        }), 500


@app.route('/api/change-password', methods=['POST'])
@login_required
def change_password():
    data = request.json
    
    if not data:
        return jsonify({'success': False, 'error': '请求体为空'}), 400
    
    old_password = data.get('old_password')
    new_password = data.get('new_password')
    
    if not old_password or not new_password:
        return jsonify({'success': False, 'error': '旧密码和新密码不能为空'}), 400
    
    auth_manager = AuthManager(DATABASE_PATH)
    result = auth_manager.change_password(request.user_name, old_password, new_password)
    
    if result['success']:
        return jsonify(result)
    else:
        return jsonify(result), 400


@app.route('/api/update-profile', methods=['POST'])
@login_required
def update_profile():
    """
    修改个人信息(手机号/部门/入职日期)
    管理员可修改任意用户，普通用户只能修改自己
    """
    data = request.json
    
    if not data:
        return jsonify({'success': False, 'error': '请求体为空'}), 400
    
    target_name = data.get('name', request.user_name)
    phone = data.get('手机号')
    dept = data.get('部门')
    hire_date = data.get('入职日期')
    
    if request.user_role != 'admin' and target_name != request.user_name:
        return jsonify({'success': False, 'error': '无权修改其他用户信息'}), 403
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT id FROM users WHERE 姓名 = ?", (target_name,))
        if not cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'error': '用户不存在'}), 404
        
        updates = []
        params = []
        
        if phone is not None:
            updates.append("手机号 = ?")
            params.append(phone)
        if dept is not None:
            updates.append("部门 = ?")
            params.append(dept)
        if hire_date is not None:
            updates.append("入职日期 = ?")
            params.append(hire_date)
        
        if not updates:
            conn.close()
            return jsonify({'success': False, 'error': '没有需要更新的字段'}), 400
        
        params.append(target_name)
        cursor.execute(f"UPDATE users SET {', '.join(updates)} WHERE 姓名 = ?", params)
        conn.commit()
        conn.close()
        
        conn2 = get_db_connection()
        cursor2 = conn2.cursor()
        cursor2.execute(
            "UPDATE summary_table SET 手机号 = COALESCE(?, 手机号), 部门 = COALESCE(?, 部门), 入职日期 = COALESCE(?, 入职日期) WHERE 姓名 = ?",
            (phone, dept, hire_date, target_name)
        )
        conn2.commit()
        conn2.close()
        
        return jsonify({'success': True, 'message': '个人信息更新成功'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/verify-token', methods=['GET'])
@login_required
def verify_token():
    return jsonify({
        'success': True,
        'name': request.user_name,
        'role': request.user_role
    })


@app.route('/api/salary', methods=['GET'])
@login_required
def get_salary():
    try:
        month = request.args.get('month', '').strip()
        salary_manager = SalaryManager(DATABASE_PATH)
        
        if month:
            available_months = salary_manager.get_available_months(request.user_name)
            if month not in available_months:
                return jsonify({
                    'success': False,
                    'error': f'{month}月份工资表尚未上传，请稍后再试'
                }), 200
            salary_data = salary_manager.get_salary_by_name_and_month(request.user_name, month)
        else:
            salary_data = salary_manager.get_salary_by_name(request.user_name)
        
        if not salary_data:
            return jsonify({
                'success': False,
                'error': '未找到工资信息'
            }), 404
        
        salary_month = month or salary_data.get('月份', '')
        sync_success = salary_manager.sync_to_summary(request.user_name, salary_month)
        
        if not sync_success:
            print(f"[WARN] 同步工资数据到汇总表失败: {request.user_name}")
        
        summary_data = salary_manager.get_summary_by_name_and_month(request.user_name, salary_month)
        
        return jsonify({
            'success': True,
            'data': {
                'salary': salary_data,
                'summary': summary_data,
                'query_count': summary_data['查询次数'] if summary_data else 0,
                'sign_status': summary_data['签收状态'] if summary_data else '未签收',
                'resign_reason': summary_data['重签原因'] if summary_data and summary_data.get('重签原因') else ''
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/available-months', methods=['GET'])
@login_required
def get_available_months():
    try:
        salary_manager = SalaryManager(DATABASE_PATH)
        months = salary_manager.get_available_months(request.user_name)
        
        return jsonify({
            'success': True,
            'months': months
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/export-my-salary', methods=['GET'])
@login_required
def export_my_salary():
    """
    导出当前用户的工资信息到CSV
    """
    try:
        import csv
        import io
        
        salary_manager = SalaryManager(DATABASE_PATH)
        salary_data = salary_manager.get_salary_by_name(request.user_name)
        summary_data = salary_manager.get_summary_by_name(request.user_name)
        
        if not salary_data:
            return jsonify({
                'success': False,
                'error': '未找到工资信息'
            }), 404
        
        output = io.StringIO()
        writer = csv.writer(output, quoting=csv.QUOTE_ALL)
        
        writer.writerow(['项目', '数值'])
        writer.writerow(['姓名', request.user_name])
        writer.writerow(['部门', salary_data.get('部门', '')])
        writer.writerow(['岗位', salary_data.get('岗位', '')])
        writer.writerow(['入职日期', salary_data.get('入职日期', '')])
        writer.writerow(['应出勤（天）', salary_data.get('应出勤天数', '')])
        writer.writerow(['实际出勤（天）', salary_data.get('实际出勤天数', '')])
        writer.writerow(['上门服务（小时）', salary_data.get('上门服务小时', '')])
        writer.writerow(['基本工资-底薪', salary_data.get('基本工资底薪', '')])
        writer.writerow(['基本工资-其它补贴', salary_data.get('基本工资其它补贴', '')])
        writer.writerow(['基本工资-合计', salary_data.get('基本工资合计', '')])
        writer.writerow(['岗位工资', salary_data.get('岗位工资', '')])
        writer.writerow(['交通费', salary_data.get('交通费', '')])
        writer.writerow(['手机费', salary_data.get('手机费', '')])
        writer.writerow(['奖金', salary_data.get('奖金', '')])
        writer.writerow(['高温费', salary_data.get('高温费', '')])
        writer.writerow(['护理员绩效工资', salary_data.get('护理员绩效工资', '')])
        writer.writerow(['应发工资', salary_data.get('应发工资', '')])
        writer.writerow(['应扣款项-缺勤扣款', salary_data.get('应扣款项缺勤扣款', '')])
        writer.writerow(['应扣款项-养老(8%)', salary_data.get('应扣款项养老', '')])
        writer.writerow(['应扣款项-医疗(2%)', salary_data.get('应扣款项医疗', '')])
        writer.writerow(['应扣款项-失业(0.5%)', salary_data.get('应扣款项失业', '')])
        writer.writerow(['应扣款项-公积金(7%)', salary_data.get('应扣款项公积金', '')])
        writer.writerow(['应扣款项-应缴个税', salary_data.get('应扣款项应缴个税', '')])
        writer.writerow(['其它扣款', salary_data.get('其它扣款', '')])
        writer.writerow(['住宿扣款', salary_data.get('住宿扣款', '')])
        writer.writerow(['水电扣款', salary_data.get('水电扣款', '')])
        writer.writerow(['实发工资', salary_data.get('实发工资', '')])
        
        if summary_data:
            writer.writerow(['查询状态', summary_data.get('查询状态', '')])
            writer.writerow(['查询次数', summary_data.get('查询次数', '')])
            writer.writerow(['签收状态', summary_data.get('签收状态', '')])
        
        output.seek(0)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"工资单_{request.user_name}_{timestamp}.csv"
        
        return send_file(
            io.BytesIO(output.getvalue().encode('utf-8-sig')),
            as_attachment=True,
            download_name=filename,
            mimetype='text/csv'
        )
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导出失败: {str(e)}'
        }), 500


@app.route('/api/signature', methods=['POST'])
@login_required
def upload_signature():
    """
    上传签名
    
    Request Body:
        {
            "signature": "base64编码的签名图片",
            "format": "PNG/JPG"
        }
    
    功能：
    1. 验证文件格式（JPG/PNG/PDF）
    2. 验证文件大小（≤10MB）
    3. 保存签名文件到本地
    4. 保存文件信息到signature_files表
    5. 更新summary_table的签收状态
    """
    try:
        data = request.json
        
        if not data:
            return jsonify({
                'success': False,
                'error': '请求体为空'
            }), 400
        
        signature_data = data.get('signature')
        file_format = data.get('format', 'PNG').upper()
        month = data.get('month', '').strip()
        
        if not signature_data:
            return jsonify({
                'success': False,
                'error': '签名数据为空'
            }), 400
        
        # 使用SignatureManager上传签名
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        result = signature_manager.upload_signature(
            request.user_name,
            signature_data,
            file_format
        )
        
        if not result['success']:
            return jsonify(result), 400
        
        # 更新签收状态
        update_success = signature_manager.update_sign_status(
            request.user_name,
            result['file_id'],
            month
        )
        
        if not update_success:
            print(f"[WARN] 更新签收状态失败: {request.user_name}")
        
        return jsonify({
            'success': True,
            'message': '签名上传成功',
            'file_id': result['file_id'],
            'file_name': result['file_name'],
            'file_size': result['file_size']
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'上传失败: {str(e)}'
        }), 500


@app.route('/api/signature/<file_id>', methods=['GET'])
@login_required
def get_signature(file_id):
    """
    获取签名文件
    
    Args:
        file_id: 文件ID
    """
    try:
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        file_info = signature_manager.get_signature_file(file_id)
        
        if not file_info:
            return jsonify({
                'success': False,
                'error': '签名文件不存在'
            }), 404
        
        file_path = file_info['存储路径']
        
        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': '签名文件不存在'
            }), 404
        
        return send_file(file_path, as_attachment=True, download_name=file_info['文件名'])
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取签名失败: {str(e)}'
        }), 500


@app.route('/api/signatures', methods=['GET'])
@login_required
def get_user_signatures():
    """
    获取当前用户的所有签名文件
    """
    try:
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        signatures = signature_manager.get_user_signatures(request.user_name)
        
        return jsonify({
            'success': True,
            'data': signatures,
            'count': len(signatures)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


# ========================================
# 管理员API
# ========================================

@app.route('/api/admin/users', methods=['GET'])
@admin_required
def get_all_users():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            """SELECT 
                u.姓名,
                u.手机号,
                u.部门,
                u.入职日期,
                u.role,
                u.last_login_time,
                u.created_at,
                s.查询状态,
                s.查询次数,
                s.签收状态,
                s.最新签收时间
               FROM users u
               LEFT JOIN summary_table s ON u.姓名 = s.姓名
               ORDER BY u.姓名"""
        )
        
        users = cursor.fetchall()
        conn.close()
        
        users_list = [dict(user) for user in users]
        
        return jsonify({
            'success': True,
            'data': users_list
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/status', methods=['GET'])
@admin_required
def get_user_status():
    """
    获取所有用户状态（管理员）
    
    返回格式：
    {
        "success": true,
        "data": [
            {
                "username": "张三",
                "viewed": true,
                "signed": false,
                "signature_path": null,
                "salary_date": "2026-04",
                "last_login": "2026-04-07 10:30:00"
            }
        ]
    }
    """
    try:
        month = request.args.get('month')
        
        if not month:
            return jsonify({
                'success': True,
                'data': [],
                'message': '请先选择月份'
            })
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            """SELECT 
                u.姓名 as username,
                st.查询状态,
                st.签收状态,
                st.签名图片 as signature_path,
                st.入职日期,
                u.last_login_time,
                s.月份
               FROM users u
               LEFT JOIN salary_table s ON u.姓名 = s.姓名 AND s.月份 = ?
               LEFT JOIN summary_table st ON u.姓名 = st.姓名 AND st.月份 = ?
               WHERE u.role = 'user'
               ORDER BY u.姓名""",
            (month, month)
        )
        
        users = cursor.fetchall()
        conn.close()
        
        users_list = []
        for user in users:
            users_list.append({
                "username": user["username"],
                "viewed": user["查询状态"] == "已查询" if user["查询状态"] else False,
                "signed": user["签收状态"] == "已签收" if user["签收状态"] else False,
                "signature_path": user["signature_path"] or "",
                "salary_date": user["入职日期"] or "",
                "last_login": user["last_login_time"] or ""
            })
        
        return jsonify({
            'success': True,
            'data': users_list
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/resign', methods=['POST'])
@admin_required
def request_resign():
    """
    要求用户重新签名（管理员）
    
    Request Body:
        {
            "target_user": "目标用户姓名",
            "reason": "重签原因（可选）"
        }
    """
    try:
        data = request.json
        
        if not data:
            return jsonify({
                'success': False,
                'error': '请求体为空'
            }), 400
        
        target_user = data.get('target_user')
        reason = data.get('reason', '签名不合格，请重新签名')
        
        if not target_user:
            return jsonify({
                'success': False,
                'error': '目标用户不能为空'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 更新汇总表签收状态和重签原因
        cursor.execute(
            """UPDATE summary_table 
               SET 签收状态 = '未签收',
                   签收时间 = NULL,
                   重签原因 = ?
               WHERE 姓名 = ?""",
            (reason, target_user)
        )
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'已要求 {target_user} 重新签名'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'操作失败: {str(e)}'
        }), 500


@app.route('/api/admin/statistics', methods=['GET'])
@admin_required
def get_statistics():
    """
    获取查询统计信息（管理员）
    """
    try:
        month = request.args.get('month')
        salary_manager = SalaryManager(DATABASE_PATH)
        stats = salary_manager.get_query_statistics(month)
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/department-statistics', methods=['GET'])
@admin_required
def get_department_statistics():
    """
    获取部门统计信息（管理员）
    """
    try:
        salary_manager = SalaryManager(DATABASE_PATH)
        stats = salary_manager.get_department_statistics()
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/summaries', methods=['GET'])
@admin_required
def get_all_summaries():
    try:
        department = request.args.get('department')
        month = request.args.get('month')
        
        salary_manager = SalaryManager(DATABASE_PATH)
        summaries = salary_manager.get_all_summaries(department, month)
        
        return jsonify({
            'success': True,
            'data': summaries,
            'count': len(summaries)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/user-salary/<name>', methods=['GET'])
@admin_required
def get_user_salary_detail(name):
    """
    获取指定用户的工资详情（管理员）
    
    Path Parameters:
        name: 用户姓名
    
    Query Parameters:
        month: 月份筛选 (格式: YYYY-MM)
    """
    try:
        month = request.args.get('month')
        salary_manager = SalaryManager(DATABASE_PATH)
        
        if month:
            salary_data = salary_manager.get_salary_by_name_and_month(name, month)
            summary_data = salary_manager.get_summary_by_name_and_month(name, month)
        else:
            salary_data = salary_manager.get_salary_by_name(name)
            summary_data = salary_manager.get_summary_by_name(name)
        
        if not salary_data:
            return jsonify({
                'success': False,
                'error': '未找到该用户的工资信息'
            }), 404
        
        return jsonify({
            'success': True,
            'data': {
                'salary': salary_data,
                'summary': summary_data
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/signatures', methods=['GET'])
@admin_required
def get_all_signatures():
    """
    获取所有签名文件（管理员）
    """
    try:
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        signatures = signature_manager.get_all_signatures()
        
        return jsonify({
            'success': True,
            'data': signatures,
            'count': len(signatures)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/signatures-with-users', methods=['GET'])
@admin_required
def get_signatures_with_users():
    """
    获取所有签名文件及用户信息（管理员）
    
    返回签名文件列表，包含：
    - 文件ID
    - 用户姓名
    - 文件格式
    - 文件大小
    - 上传时间
    - 预览URL
    - 签收状态
    - 重签原因
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 查询所有签名文件和用户信息
        cursor.execute(
            """SELECT 
                sf.文件ID,
                sf.用户姓名,
                sf.文件格式,
                sf.文件大小,
                sf.上传时间,
                sf.存储路径,
                st.签收状态,
                st.重签原因,
                st.部门,
                st.实发工资
            FROM signature_files sf
            LEFT JOIN summary_table st ON sf.用户姓名 = st.姓名
            ORDER BY sf.上传时间 DESC"""
        )
        
        signatures = []
        for row in cursor.fetchall():
            signatures.append({
                '文件ID': row[0],
                '用户姓名': row[1],
                '文件格式': row[2],
                '文件大小': row[3],
                '上传时间': row[4],
                '存储路径': row[5],
                '签收状态': row[6] or '未签收',
                '重签原因': row[7] or '',
                '部门': row[8] or '',
                '实发工资': row[9] or 0,
                '预览URL': f'/api/admin/signature-preview/{row[0]}'
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'data': signatures,
            'count': len(signatures)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/signature-preview/<file_id>', methods=['GET'])
@admin_required
def get_signature_preview(file_id):
    """
    获取签名文件预览（管理员）
    
    返回签名文件图片，用于在浏览器中预览
    """
    try:
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        file_info = signature_manager.get_signature_file(file_id)
        
        if not file_info:
            return jsonify({
                'success': False,
                'error': '签名文件不存在'
            }), 404
        
        file_path = file_info['存储路径']
        
        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': '签名文件不存在'
            }), 404
        
        # 返回文件用于预览（不作为附件下载）
        return send_file(file_path, as_attachment=False, download_name=file_info['文件名'])
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取签名失败: {str(e)}'
        }), 500


@app.route('/api/admin/signature/<username>', methods=['GET'])
@admin_required
def get_user_signature(username):
    """
    获取指定用户的签名图片（管理员）
    
    Query Parameters:
        month: 月份筛选 (格式: YYYY-MM)
    
    直接返回图片数据，用于在APP中显示
    """
    try:
        month = request.args.get('month')
        conn = get_db_connection()
        cursor = conn.cursor()
        
        if month:
            cursor.execute(
                "SELECT 签名图片 FROM summary_table WHERE 姓名 = ? AND 月份 = ?",
                (username, month)
            )
        else:
            cursor.execute(
                "SELECT 签名图片 FROM summary_table WHERE 姓名 = ? ORDER BY updated_at DESC LIMIT 1",
                (username,)
            )
        
        result = cursor.fetchone()
        conn.close()
        
        if not result or not result['签名图片']:
            return jsonify({
                'success': False,
                'error': '该用户暂无签名'
            }), 404
        
        signature_name = result['签名图片']
        
        possible_paths = [
            os.path.join(SIGNATURE_DIR, f"{signature_name}.png"),
            os.path.join(SIGNATURE_DIR, f"{signature_name}.jpg"),
            os.path.join(SIGNATURE_DIR, signature_name),
        ]
        
        signature_path = None
        for path in possible_paths:
            if os.path.exists(path):
                signature_path = path
                break
        
        if not signature_path:
            return jsonify({
                'success': False,
                'error': f'签名文件不存在: {signature_name}'
            }), 404
        
        return send_file(signature_path, as_attachment=False, download_name=f'{username}_signature.png')
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取签名失败: {str(e)}'
        }), 500


@app.route('/api/admin/signatures/download', methods=['GET'])
@admin_required
def download_all_signatures():
    """
    下载所有签名文件（ZIP格式）
    """
    try:
        import zipfile
        import io
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            "SELECT 姓名, 签名图片 FROM summary_table WHERE 签名图片 IS NOT NULL AND 签名图片 != ''"
        )
        
        signatures = cursor.fetchall()
        conn.close()
        
        if not signatures:
            return jsonify({
                'success': False,
                'error': '暂无签名文件'
            }), 404
        
        memory_file = io.BytesIO()
        file_count = 0
        
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for sig in signatures:
                username = sig['姓名']
                signature_name = sig['签名图片']
                
                possible_paths = [
                    os.path.join(SIGNATURE_DIR, f"{signature_name}.png"),
                    os.path.join(SIGNATURE_DIR, f"{signature_name}.jpg"),
                    os.path.join(SIGNATURE_DIR, signature_name),
                ]
                
                for path in possible_paths:
                    if os.path.exists(path):
                        ext = os.path.splitext(path)[1]
                        zf.write(path, f'{username}_signature{ext}')
                        file_count += 1
                        break
        
        if file_count == 0:
            return jsonify({
                'success': False,
                'error': '签名文件均不存在'
            }), 404
        
        memory_file.seek(0)
        
        return send_file(
            memory_file,
            as_attachment=True,
            download_name='all_signatures.zip',
            mimetype='application/zip'
        )
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'下载失败: {str(e)}'
        }), 500


@app.route('/api/admin/signature-statistics', methods=['GET'])
@admin_required
def get_signature_statistics():
    """
    获取签名文件统计信息（管理员）
    """
    try:
        signature_manager = SignatureManager(DATABASE_PATH, SIGNATURE_DIR)
        
        stats = signature_manager.get_signature_statistics()
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/sync-all', methods=['POST'])
@admin_required
def sync_all_to_summary():
    """
    同步所有工资数据到汇总表（管理员）
    """
    try:
        data_manager = DataManager(DATABASE_PATH)
        
        result = data_manager.sync_all_to_summary()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'同步失败: {str(e)}'
        }), 500


@app.route('/api/admin/import-csv', methods=['POST'])
@admin_required
def import_salary_from_csv():
    """
    从CSV文件导入工资数据（管理员）
    
    Request Body:
        {
            "file_path": "CSV文件路径"
        }
    """
    try:
        data = request.json
        
        if not data or 'file_path' not in data:
            return jsonify({
                'success': False,
                'error': '缺少file_path参数'
            }), 400
        
        file_path = data['file_path']
        
        data_manager = DataManager(DATABASE_PATH)
        
        result = data_manager.import_salary_from_csv(file_path)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导入失败: {str(e)}'
        }), 500


@app.route('/api/admin/export-csv', methods=['GET'])
@admin_required
def export_summary_to_csv():
    try:
        department = request.args.get('department')
        month = request.args.get('month')
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        month_suffix = f"_{month}" if month else ""
        filename = f"summary{month_suffix}_{timestamp}.xlsx"
        output_path = os.path.join(os.path.dirname(__file__), 'exports', filename)
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        data_manager = DataManager(DATABASE_PATH)
        
        result = data_manager.export_summary_to_excel(output_path, department, month)
        
        if result['success']:
            return send_file(
                output_path,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        else:
            return jsonify(result), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导出失败: {str(e)}'
        }), 500


@app.route('/api/admin/available-months', methods=['GET'])
@admin_required
def get_admin_available_months():
    try:
        salary_manager = SalaryManager(DATABASE_PATH)
        months = salary_manager.get_available_months()
        
        return jsonify({
            'success': True,
            'months': months
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/device-bindings', methods=['GET'])
@admin_required
def get_device_bindings():
    try:
        device_manager = DeviceManager(DATABASE_PATH)
        bindings = device_manager.get_all_bindings()
        
        return jsonify({
            'success': True,
            'bindings': bindings
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/unbind-device', methods=['POST'])
@admin_required
def unbind_device():
    try:
        data = request.json
        username = data.get('username')
        
        if not username:
            return jsonify({
                'success': False,
                'error': '用户名不能为空'
            }), 400
        
        device_manager = DeviceManager(DATABASE_PATH)
        success = device_manager.unbind_device(username)
        
        if success:
            return jsonify({
                'success': True,
                'message': f'已解绑用户 {username} 的设备'
            })
        else:
            return jsonify({
                'success': False,
                'error': '解绑失败'
            }), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'解绑失败: {str(e)}'
        }), 500


@app.route('/api/admin/failed-attempts', methods=['GET'])
@admin_required
def get_failed_attempts():
    try:
        limit = request.args.get('limit', 50, type=int)
        device_manager = DeviceManager(DATABASE_PATH)
        attempts = device_manager.get_failed_attempts(limit)
        
        return jsonify({
            'success': True,
            'attempts': attempts
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/export-report', methods=['GET'])
@admin_required
def export_statistics_report():
    """
    导出统计报表（管理员）- Excel格式
    
    Query Parameters:
        type: 报表类型 (summary/unsigned/queries/full)
        month: 筛选月份（可选）
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
        
        export_type = request.args.get('type', 'summary')
        month = request.args.get('month', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        wb = openpyxl.Workbook()
        ws = wb.active
        
        title_font = Font(bold=True, size=14)
        header_font = Font(bold=True, size=11)
        header_fill = PatternFill(start_color="1E88E5", end_color="1E88E5", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        month_suffix = f"_{month}" if month else ""
        filename = f"salary_report_{export_type}{month_suffix}_{timestamp}.xlsx"
        
        if export_type == 'summary':
            ws.title = "签收统计报表"
            ws['A1'] = "人事管理系统 - 签收统计报表"
            ws['A1'].font = title_font
            ws.merge_cells('A1:E1')
            
            ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            if month:
                ws['A3'] = f"筛选月份: {month}"
            
            cursor.execute("SELECT COUNT(*) FROM users WHERE role = 'user'")
            total_users = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM summary_table WHERE 签收状态 = '已签收'")
            signed_count = cursor.fetchone()[0]
            
            row = 5
            ws[f'A{row}'] = "统计概览"
            ws[f'A{row}'].font = header_font
            row += 1
            ws[f'A{row}'] = "总用户数"
            ws[f'B{row}'] = total_users
            row += 1
            ws[f'A{row}'] = "已签收"
            ws[f'B{row}'] = signed_count
            row += 1
            ws[f'A{row}'] = "未签收"
            ws[f'B{row}'] = total_users - signed_count
            row += 2
            
            headers = ['姓名', '签收状态', '签收时间', '查询状态', '部门']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center')
            
            row += 1
            query = """
                SELECT u.姓名, s.签收状态, s.签收时间, s.查询状态, u.部门
                FROM users u
                LEFT JOIN summary_table s ON u.姓名 = s.姓名
                WHERE u.role = 'user'
                ORDER BY u.姓名
            """
            cursor.execute(query)
            users = cursor.fetchall()
            
            for user in users:
                ws.cell(row=row, column=1, value=user['姓名']).border = thin_border
                ws.cell(row=row, column=2, value=user['签收状态'] or '未签收').border = thin_border
                ws.cell(row=row, column=3, value=user['签收时间'] or '-').border = thin_border
                ws.cell(row=row, column=4, value=user['查询状态'] or '未查询').border = thin_border
                ws.cell(row=row, column=5, value=user['部门'] or '-').border = thin_border
                row += 1
                
        elif export_type == 'unsigned':
            ws.title = "未签收用户列表"
            ws['A1'] = "人事管理系统 - 未签收用户列表"
            ws['A1'].font = title_font
            ws.merge_cells('A1:D1')
            
            ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            if month:
                ws['A3'] = f"筛选月份: {month}"
            
            row = 5
            headers = ['姓名', '部门', '手机号', '入职日期']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center')
            
            row += 1
            query = """
                SELECT u.姓名, u.部门, u.手机号, u.入职日期
                FROM users u
                LEFT JOIN summary_table s ON u.姓名 = s.姓名
                WHERE u.role = 'user' AND (s.签收状态 IS NULL OR s.签收状态 != '已签收')
                ORDER BY u.姓名
            """
            cursor.execute(query)
            users = cursor.fetchall()
            
            for user in users:
                ws.cell(row=row, column=1, value=user['姓名']).border = thin_border
                ws.cell(row=row, column=2, value=user['部门'] or '-').border = thin_border
                ws.cell(row=row, column=3, value=user['手机号'] or '-').border = thin_border
                ws.cell(row=row, column=4, value=user['入职日期'] or '-').border = thin_border
                row += 1
                
        elif export_type == 'queries':
            ws.title = "异议记录报表"
            ws['A1'] = "人事管理系统 - 异议记录报表"
            ws['A1'].font = title_font
            ws.merge_cells('A1:D1')
            
            ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            row = 4
            headers = ['姓名', '异议内容', '提交时间', '部门']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center')
            
            row += 1
            query = """
                SELECT s.姓名, s.异议内容, s.查询时间, u.部门
                FROM summary_table s
                LEFT JOIN users u ON s.姓名 = u.姓名
                WHERE s.查询状态 = '有异议' AND s.异议内容 IS NOT NULL
                ORDER BY s.查询时间 DESC
            """
            cursor.execute(query)
            queries = cursor.fetchall()
            
            for q in queries:
                ws.cell(row=row, column=1, value=q['姓名']).border = thin_border
                ws.cell(row=row, column=2, value=q['异议内容'] or '-').border = thin_border
                ws.cell(row=row, column=3, value=q['查询时间'] or '-').border = thin_border
                ws.cell(row=row, column=4, value=q['部门'] or '-').border = thin_border
                row += 1
                
        elif export_type == 'full':
            if not month:
                conn.close()
                return jsonify({
                    'success': False,
                    'error': '请选择月份后再导出完整签收记录'
                }), 400
            
            cursor.execute(
                "SELECT COUNT(*) FROM salary_table WHERE 月份 = ?",
                (month,)
            )
            month_count = cursor.fetchone()[0]
            
            if month_count == 0:
                conn.close()
                return jsonify({
                    'success': False,
                    'error': f'所选月份 {month} 的工资表尚未上传，请先上传工资表'
                }), 400
            
            conn.close()
            
            month_display = month.replace('-', '年') + '月'
            filename = f"{month_display}签收记录表.xlsx"
            
            data_manager = DataManager(DATABASE_PATH)
            output_path = os.path.join(os.path.dirname(__file__), 'exports', filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            result = data_manager.export_summary_to_excel(output_path, month=month)
            
            if result['success']:
                if result['exported_count'] == 0:
                    return jsonify({
                        'success': False,
                        'error': '没有可导出的数据，请先上传工资表'
                    }), 400
                    
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )
            else:
                return jsonify({
                    'success': False,
                    'error': '导出失败'
                }), 500
        
        for col in range(1, ws.max_column + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
        
        conn.close()
        
        output_path = os.path.join(os.path.dirname(__file__), 'exports', filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        wb.save(output_path)
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导出失败: {str(e)}'
        }), 500


@app.route('/api/admin/data-summary', methods=['GET'])
@admin_required
def get_data_summary():
    """
    获取数据汇总信息（管理员）
    """
    try:
        data_manager = DataManager(DATABASE_PATH)
        
        summary = data_manager.get_data_summary()
        
        return jsonify({
            'success': True,
            'data': summary
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


@app.route('/api/admin/check-month-data', methods=['GET'])
@admin_required
def check_month_data():
    """
    检查指定月份是否已有工资数据（管理员）
    """
    try:
        month = request.args.get('month', '').strip()
        if not month:
            return jsonify({
                'success': False,
                'error': '请提供月份参数'
            }), 400
        
        data_manager = DataManager(DATABASE_PATH)
        conn = data_manager.get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            "SELECT COUNT(*) as count FROM salary_table WHERE 月份 = ?",
            (month,)
        )
        
        result = cursor.fetchone()
        count = result['count'] if result else 0
        
        conn.close()
        
        return jsonify({
            'success': True,
            'has_data': count > 0,
            'count': count
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'查询失败: {str(e)}'
        }), 500


def clear_month_data(month: str) -> dict:
    """
    清理指定月份的所有数据（用于重复上传时重置）
    
    包括：
    1. 删除 salary_table 中该月份的工资数据
    2. 删除 summary_table 中该月份的签收记录
    3. 删除 signature_files 中该月份相关的签名文件
    
    Args:
        month: 月份（如 2024-01）
    
    Returns:
        dict: 清理结果
    """
    result = {
        'salary_deleted': 0,
        'summary_deleted': 0,
        'signatures_deleted': 0,
        'signature_files_deleted': 0
    }
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("DELETE FROM salary_table WHERE 月份 = ?", (month,))
        result['salary_deleted'] = cursor.rowcount
        
        cursor.execute("SELECT 签名图片 FROM summary_table WHERE 月份 = ? AND 签名图片 IS NOT NULL AND 签名图片 != ''", (month,))
        signature_ids = [row['签名图片'] for row in cursor.fetchall()]
        
        cursor.execute("DELETE FROM summary_table WHERE 月份 = ?", (month,))
        result['summary_deleted'] = cursor.rowcount
        
        if signature_ids:
            for sig_id in signature_ids:
                cursor.execute("SELECT 存储路径 FROM signature_files WHERE 文件ID = ?", (sig_id,))
                sig_row = cursor.fetchone()
                if sig_row:
                    file_path = sig_row['存储路径']
                    if os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            result['signature_files_deleted'] += 1
                        except Exception as e:
                            print(f"[WARN] 删除签名文件失败: {file_path}, {e}")
            
            placeholders = ','.join(['?' for _ in signature_ids])
            cursor.execute(f"DELETE FROM signature_files WHERE 文件ID IN ({placeholders})", signature_ids)
            result['signatures_deleted'] = cursor.rowcount
        
        conn.commit()
        conn.close()
        
        print(f"[INFO] 清理月份 {month} 数据: 工资记录={result['salary_deleted']}, 签收记录={result['summary_deleted']}, 签名文件={result['signatures_deleted']}")
        
        return result
        
    except Exception as e:
        print(f"[ERROR] 清理月份数据失败: {e}")
        return result


@app.route('/api/admin/upload-salary', methods=['POST'])
@admin_required
def upload_salary_file():
    """
    上传工资表文件（管理员）
    
    支持 Excel (.xlsx, .xls) 和 CSV (.csv) 格式
    文件大小限制：10MB
    
    如果同月份已有数据，将自动清理旧数据并重置签收状态
    """
    try:
        month = request.form.get('month', '').strip()
        if not month:
            return jsonify({
                'success': False,
                'error': '请选择工资表所属月份'
            }), 400
        
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': '没有上传文件'
            }), 400
        
        file = request.files['file']
        
        # 获取原始文件名
        original_filename = file.filename
        
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': '没有选择文件'
            }), 400
        
        # 检查文件格式 - 多种方式获取扩展名
        allowed_extensions = {'xlsx', 'xls', 'csv'}
        file_ext = None
        
        # 方法1: 从文件名获取
        if original_filename and '.' in original_filename:
            _, ext = os.path.splitext(original_filename)
            if ext:
                file_ext = ext.lower().lstrip('.')
        
        # 方法2: 从Content-Type推断
        if not file_ext:
            content_type = request.content_type or ''
            
            if 'excel' in content_type.lower() or 'spreadsheet' in content_type.lower():
                file_ext = 'xlsx'
            elif 'csv' in content_type.lower():
                file_ext = 'csv'
        
        # 方法3: 读取文件头判断
        if not file_ext:
            file.seek(0)
            header = file.read(8)
            file.seek(0)
            
            # Excel xlsx 是 ZIP 格式，以 PK 开头
            if header[:4] == b'PK\x03\x04':
                file_ext = 'xlsx'
            # Excel xls 以 D0 CF 11 E0 开头 (OLE格式)
            elif header[:4] == b'\xd0\xcf\x11\xe0':
                file_ext = 'xls'
        
        if not file_ext:
            return jsonify({
                'success': False,
                'error': f'无法识别文件格式。请确保文件是 Excel (.xlsx, .xls) 或 CSV (.csv) 格式'
            }), 400
        
        if file_ext not in allowed_extensions:
            return jsonify({
                'success': False,
                'error': f'不支持的文件格式: .{file_ext}。支持的格式: .xlsx, .xls, .csv'
            }), 400
        
        # 使用UUID生成唯一文件名，避免中文文件名问题
        import uuid
        filename = f"salary_{uuid.uuid4().hex[:8]}.{file_ext}"
        
        # 检查文件大小
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({
                'success': False,
                'error': f'文件大小超出限制: {file_size}字节。最大允许: 10MB'
            }), 400
        
        # 保存文件到临时目录
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        temp_file_path = os.path.join(temp_dir, filename)
        file.save(temp_file_path)
        
        clear_result = clear_month_data(month)
        
        if clear_result['salary_deleted'] > 0:
            print(f"[INFO] 检测到同月份数据，已清理: 工资记录 {clear_result['salary_deleted']} 条, 签收记录 {clear_result['summary_deleted']} 条, 签名文件 {clear_result['signatures_deleted']} 个")
        
        # 根据文件格式处理
        data_manager = DataManager(DATABASE_PATH)
        
        if file_ext == 'csv':
            result = data_manager.import_salary_from_csv(temp_file_path, month)
        else:
            result = import_salary_from_excel(temp_file_path, data_manager, month)
        
        if clear_result['salary_deleted'] > 0:
            if 'message' not in result:
                result['message'] = ''
            result['message'] = f"检测到同月份数据，已清理旧数据并重置签收状态。{result['message']}"
            result['cleared'] = clear_result
        
        # 删除临时文件
        try:
            os.remove(temp_file_path)
        except:
            pass
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'上传失败: {str(e)}'
        }), 500


def import_salary_from_excel(file_path: str, data_manager, month: str = '') -> dict:
    """
    从Excel文件导入工资数据
    
    表头规则：
    - A-I列（索引0-8）：表头在第一行（A1-I1）
    - J-AA列（索引9-26）：表头在第二行（J2-AA2）
    - AB-AD列（索引27-29）：表头在第一行（AB1-AD1）
    
    Args:
        file_path: Excel文件路径
        data_manager: DataManager实例
    
    Returns:
        dict: 导入结果
    """
    try:
        import openpyxl
        
        wb_data_only = openpyxl.load_workbook(file_path, data_only=True)
        ws_data_only = wb_data_only.active
        
        wb_formulas = openpyxl.load_workbook(file_path, data_only=False)
        ws_formulas = wb_formulas.active
        
        ws = ws_data_only
        ws_formula = ws_formulas
        
        imported_count = 0
        error_count = 0
        errors = []
        
        FIELD_MAPPING = {
            '序号': ['序号'],
            '部门': ['部门'],
            '姓名': ['姓名'],
            '入职日期': ['入职日期', '入职 日期', '入职\n日期'],
            '是否代扣社保': ['是否代扣社保', '是否代扣\n社保', '是否\n代扣社保'],
            '岗位': ['岗位'],
            '应出勤天数': ['应出勤（天）', '应出勤天数', '应出勤\n（天）'],
            '实际出勤天数': ['实际出勤（天）', '实际出勤天数', '实际出勤\n（天）'],
            '上门服务小时': ['上门服务（小时）', '上门服务小时', '上门服务\n（小时）'],
            '基本工资底薪': ['基本工资-底薪', '基本工资底薪'],
            '基本工资合计': ['基本工资-基本工资', '基本工资合计', '基本工资'],
            '岗位工资': ['岗位工资', '岗位\n工资'],
            '交通费': ['交通费', '交通\n费'],
            '手机费': ['手机费', '手机\n费'],
            '奖金': ['奖金'],
            '高温费': ['高温费', '高温\n费'],
            '应扣款项缺勤扣款': ['应扣款项-缺勤扣款', '应扣款项缺勤扣款', '缺勤扣款'],
            '应扣款项养老': ['应扣款项-养老(8%)', '养老(8%)', '养老\n(8%)', '应扣款项养老'],
            '应扣款项医疗': ['应扣款项-医疗(2%)', '医疗(2%)', '医疗\n(2%)', '应扣款项医疗'],
            '应扣款项失业': ['应扣款项-失业(0.5%)', '失业(0.5%)', '失业\n(0.5%)', '应扣款项失业'],
            '应扣款项公积金': ['应扣款项-公积金(7%)', '公积金(7%)', '公积金\n(7%)', '应扣款项公积金'],
            '应扣款项应缴个税': ['应扣款项-应缴个税', '应缴个税', '应缴\n个税', '应扣款项应缴个税'],
            '其它扣款': ['其它扣款', '其他扣款', '其它\n扣款', '其它捐款', '其他捐款'],
            '住宿扣款': ['住宿扣款', '住宿\n扣款', '住宿捐款'],
            '水电扣款': ['水电扣款', '水电\n扣款', '水电捐款'],
            '实发工资': ['实发工资', '实发\n工资'],
            '护理员绩效工资': ['护理员绩效工资', '护理员\n绩效工资'],
            '应发工资': ['应发工资', '应发\n工资'],
            '基本工资其它补贴': ['基本工资-其它补贴', '基本工资-其他补贴', '基本工资其它补贴', '其它补贴', '其他补贴'],
        }
        
        def normalize_header(header):
            if not header:
                return ''
            return header.replace('\n', '').replace(' ', '').replace('（', '(').replace('）', ')').strip()
        
        def find_db_field(excel_header):
            normalized = normalize_header(excel_header)
            for db_field, excel_fields in FIELD_MAPPING.items():
                for ef in excel_fields:
                    if normalize_header(ef) == normalized:
                        return db_field
            return None
        
        def col_letter_to_index(col_letter):
            result = 0
            for char in col_letter.upper():
                result = result * 26 + (ord(char) - ord('A') + 1)
            return result - 1
        
        J_COL_INDEX = col_letter_to_index('J')
        AA_COL_INDEX = col_letter_to_index('AA')
        AB_COL_INDEX = col_letter_to_index('AB')
        AD_COL_INDEX = col_letter_to_index('AD')
        
        row1 = list(ws.iter_rows(min_row=1, max_row=1))[0]
        row2 = list(ws.iter_rows(min_row=2, max_row=2))[0]
        
        total_cols = len(row1)
        headers = []
        skip_columns = set()
        
        for col_idx in range(total_cols):
            if col_idx < J_COL_INDEX:
                header_val1 = str(row1[col_idx].value) if row1[col_idx].value else ''
                header_val2 = str(row2[col_idx].value) if row2[col_idx].value else ''
                header_val = header_val1 if header_val1 and header_val1 != 'None' else header_val2
            elif col_idx < AA_COL_INDEX:
                header_val = str(row2[col_idx].value) if row2[col_idx].value else ''
            else:
                header_val1 = str(row1[col_idx].value) if row1[col_idx].value else ''
                header_val2 = str(row2[col_idx].value) if row2[col_idx].value else ''
                header_val = header_val1 if header_val1 and header_val1 != 'None' else header_val2
            
            db_field = find_db_field(header_val)
            
            if not header_val or header_val == 'None' or header_val.strip() == '':
                skip_columns.add(col_idx)
                headers.append(None)
            else:
                headers.append(db_field if db_field else normalize_header(header_val))
        
        name_col_idx = None
        for idx, h in enumerate(headers):
            if h == '姓名':
                name_col_idx = idx
                break
        
        if name_col_idx is None:
            return {
                'success': False,
                'error': 'Excel文件缺少"姓名"列'
            }
        
        data_start_row = 3 if any(row2[col_idx].value for col_idx in range(J_COL_INDEX, min(AA_COL_INDEX + 1, total_cols))) else 2
        
        conn = data_manager.get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名 FROM users")
        existing_users = set(row['姓名'] for row in cursor.fetchall())
        
        new_users = set()
        
        for row_num in range(3, ws.max_row + 1):
            row = list(ws.iter_rows(min_row=row_num, max_row=row_num))[0]
            try:
                name = str(row[name_col_idx].value).strip() if row[name_col_idx].value else ''
                
                if not name or name == 'None':
                    continue
                
                if name not in existing_users:
                    new_users.add(name)
                
                row_data = {}
                for idx, header in enumerate(headers):
                    if header and idx < len(row):
                        value = row[idx].value
                        if value is not None:
                            row_data[header] = value
                
                cursor.execute(
                    "SELECT id FROM salary_table WHERE 姓名 = ? AND 月份 = ?",
                    (name, month)
                )
                
                existing = cursor.fetchone()
                
                valid_fields = ['序号', '部门', '入职日期', '是否代扣社保', '岗位',
                               '应出勤天数', '实际出勤天数', '上门服务小时',
                               '基本工资底薪', '基本工资其它补贴', '基本工资合计',
                               '岗位工资', '交通费', '手机费', '奖金', '高温费',
                               '应扣款项缺勤扣款', '应扣款项养老', '应扣款项医疗',
                               '应扣款项失业', '应扣款项公积金', '应扣款项应缴个税',
                               '其它扣款', '住宿扣款', '水电扣款',
                               '实发工资', '护理员绩效工资', '应发工资']
                
                if existing:
                    update_fields = []
                    update_values = []
                    
                    for field, value in row_data.items():
                        if field != '姓名' and field in valid_fields:
                            update_fields.append(f"{field} = ?")
                            update_values.append(value)
                    
                    if update_fields:
                        update_fields.append("updated_at = ?")
                        update_values.append(datetime.now())
                        update_values.append(name)
                        update_values.append(month)
                        
                        sql = f"UPDATE salary_table SET {', '.join(update_fields)} WHERE 姓名 = ? AND 月份 = ?"
                        cursor.execute(sql, update_values)
                else:
                    fields = ['姓名', '月份']
                    placeholders = ['?', '?']
                    values = [name, month]
                    
                    for field, value in row_data.items():
                        if field != '姓名' and field in valid_fields:
                            fields.append(field)
                            placeholders.append('?')
                            values.append(value)
                    
                    sql = f"INSERT INTO salary_table ({', '.join(fields)}) VALUES ({', '.join(placeholders)})"
                    cursor.execute(sql, values)

                imported_count += 1
                
            except Exception as e:
                errors.append(f"第{row_num}行: {str(e)}")
                error_count += 1
        
        conn.commit()
        conn.close()
        
        print(f"[INFO] Excel导入完成: 成功 {imported_count} 条, 失败 {error_count} 条")
        
        result = {
            'success': True,
            'imported_count': imported_count,
            'error_count': error_count,
            'errors': errors[:10]
        }
        
        if new_users:
            result['new_users'] = list(new_users)
            result['message'] = f'导入成功 {imported_count} 条记录。发现 {len(new_users)} 个新用户：{", ".join(sorted(new_users))}。这些用户的工资信息暂不显示，请前往"人员管理"页面创建账号后即可显示。'
        
        return result
        
    except Exception as e:
        print(f"[ERROR] Excel导入失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'imported_count': 0,
            'error_count': 0,
            'errors': [str(e)]
        }


@app.route('/api/admin/users/add', methods=['POST'])
@admin_required
def add_user():
    """
    添加用户（管理员）
    
    请求体：
    {
        "username": "用户名",
        "password": "密码",
        "phone": "手机号",
        "role": "user" 或 "admin"
    }
    """
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '')
        phone = data.get('phone', '').strip()
        role = data.get('role', 'user')
        
        if not username:
            return jsonify({
                'success': False,
                'error': '用户名不能为空'
            }), 400
        
        if not password:
            return jsonify({
                'success': False,
                'error': '密码不能为空'
            }), 400
        
        if role not in ['user', 'admin']:
            return jsonify({
                'success': False,
                'error': '角色必须是 user 或 admin'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名 FROM users WHERE 姓名 = ?", (username,))
        if cursor.fetchone():
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户名已存在'
            }), 400
        
        from werkzeug.security import generate_password_hash
        hashed_password = generate_password_hash(password, method='pbkdf2:sha256:10000')
        
        cursor.execute(
            "INSERT INTO users (姓名, 密码, 手机号, role) VALUES (?, ?, ?, ?)",
            (username, hashed_password, phone if phone else None, role)
        )
        
        if role == 'user':
            cursor.execute(
                """INSERT INTO summary_table (序号, 部门, 姓名, 入职日期, 手机号, 签收状态, 查询状态)
                VALUES (?, ?, ?, date('now'), ?, '未签收', '未查询')""",
                (f'NEW-{username[:4]}', '待分配', username, phone if phone else None)
            )
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 添加成功'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'添加用户失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/batch-add', methods=['POST'])
@admin_required
def batch_add_users():
    """
    批量添加用户（管理员）
    
    请求体：
    {
        "users": [
            {"name": "张三", "department": "技术部", "phone": "13800138000"},
            {"name": "李四", "department": "财务部", "phone": "13900139000"}
        ]
    }
    
    密码自动设置为手机号
    """
    try:
        data = request.get_json()
        users = data.get('users', [])
        
        if not users or not isinstance(users, list):
            return jsonify({
                'success': False,
                'error': '请提供用户列表'
            }), 400
        
        results = {
            'success_count': 0,
            'skip_count': 0,
            'error_count': 0,
            'details': []
        }
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        for user in users:
            name = user.get('name', '').strip()
            department = user.get('department', '').strip()
            phone = user.get('phone', '').strip()
            
            if not name or not department or not phone:
                results['skip_count'] += 1
                results['details'].append(f'{name or "未知"}: 信息不完整，跳过')
                continue
            
            cursor.execute("SELECT 姓名 FROM users WHERE 姓名 = ?", (name,))
            if cursor.fetchone():
                results['skip_count'] += 1
                results['details'].append(f'{name}: 用户已存在，跳过')
                continue
            
            from werkzeug.security import generate_password_hash
            hashed_password = generate_password_hash(phone, method='pbkdf2:sha256:10000')
            
            cursor.execute(
                "INSERT INTO users (姓名, 密码, 手机号, 部门, role) VALUES (?, ?, ?, ?, 'user')",
                (name, hashed_password, phone, department)
            )
            
            cursor.execute(
                """INSERT INTO summary_table (序号, 部门, 姓名, 入职日期, 手机号, 签收状态, 查询状态)
                VALUES (?, ?, ?, date('now'), ?, '未签收', '未查询')""",
                (f'NEW-{name[:4]}', department, name, phone)
            )
            
            results['success_count'] += 1
            results['details'].append(f'{name}: 添加成功')
        
        conn.commit()
        conn.close()
        
        message = f'批量导入完成！成功 {results["success_count"]} 人'
        if results['skip_count'] > 0:
            message += f'，跳过 {results["skip_count"]} 人'
        message += '\n\n详情：\n' + '\n'.join(results['details'])
        
        return jsonify({
            'success': True,
            'message': message,
            'data': results
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'批量添加用户失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/delete', methods=['POST'])
@admin_required
def delete_user_by_post():
    """
    删除用户（管理员）- POST方式
    
    Request Body:
    {
        "username": "要删除的用户名"
    }
    """
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        
        if not username:
            return jsonify({
                'success': False,
                'error': '请提供用户名'
            }), 400
        
        if username == 'admin':
            return jsonify({
                'success': False,
                'error': '不能删除admin账户'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名, role FROM users WHERE 姓名 = ?", (username,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户不存在'
            }), 404
        
        cursor.execute("DELETE FROM signature_files WHERE 用户姓名 = ?", (username,))
        cursor.execute("DELETE FROM summary_table WHERE 姓名 = ?", (username,))
        cursor.execute("DELETE FROM users WHERE 姓名 = ?", (username,))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 已删除'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'删除用户失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/reset-password', methods=['POST'])
@admin_required
def reset_user_password():
    """
    重置用户密码（管理员）
    
    Request Body:
    {
        "username": "用户名",
        "new_password": "新密码"
    }
    """
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        new_password = data.get('new_password', '').strip()
        
        if not username:
            return jsonify({
                'success': False,
                'error': '请提供用户名'
            }), 400
        
        if not new_password:
            return jsonify({
                'success': False,
                'error': '请提供新密码'
            }), 400
        
        if len(new_password) < 6:
            return jsonify({
                'success': False,
                'error': '密码长度至少6位'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名 FROM users WHERE 姓名 = ?", (username,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户不存在'
            }), 404
        
        from werkzeug.security import generate_password_hash
        hashed_password = generate_password_hash(new_password, method='pbkdf2:sha256:10000')
        
        cursor.execute(
            "UPDATE users SET 密码 = ? WHERE 姓名 = ?",
            (hashed_password, username)
        )
        
        cursor.execute(
            "UPDATE users SET token = NULL, token_expire_time = NULL WHERE 姓名 = ?",
            (username,)
        )
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 密码已重置'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'重置密码失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/update-profile', methods=['POST'])
@admin_required
def update_user_profile():
    """
    更新用户信息（管理员）
    
    Request Body:
    {
        "username": "用户名",
        "phone": "手机号",
        "department": "部门"
    }
    """
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        phone = data.get('phone', '').strip() if data.get('phone') else None
        department = data.get('department', '').strip() if data.get('department') else None
        
        if not username:
            return jsonify({
                'success': False,
                'error': '请提供用户名'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名 FROM users WHERE 姓名 = ?", (username,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户不存在'
            }), 404
        
        updates = []
        params = []
        
        if phone is not None:
            updates.append("手机号 = ?")
            params.append(phone)
        if department is not None:
            updates.append("部门 = ?")
            params.append(department)
        
        if updates:
            params.append(username)
            cursor.execute(f"UPDATE users SET {', '.join(updates)} WHERE 姓名 = ?", params)
            
            if department is not None:
                cursor.execute("UPDATE summary_table SET 部门 = ? WHERE 姓名 = ?", (department, username))
            
            conn.commit()
        
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 信息已更新'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'更新用户信息失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/<username>', methods=['DELETE'])
@admin_required
def delete_user(username):
    """
    删除用户（管理员）
    
    注意：不能删除admin账户
    """
    try:
        if username == 'admin':
            return jsonify({
                'success': False,
                'error': '不能删除admin账户'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名, role FROM users WHERE 姓名 = ?", (username,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户不存在'
            }), 404
        
        cursor.execute("DELETE FROM signature_files WHERE 用户姓名 = ?", (username,))
        cursor.execute("DELETE FROM summary_table WHERE 姓名 = ?", (username,))
        cursor.execute("DELETE FROM users WHERE 姓名 = ?", (username,))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 已删除'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'删除用户失败: {str(e)}'
        }), 500


@app.route('/api/admin/users/<username>/password', methods=['PUT'])
@admin_required
def update_user_password(username):
    """
    修改用户密码（管理员）
    
    请求体：
    {
        "new_password": "新密码"
    }
    """
    try:
        data = request.get_json()
        new_password = data.get('new_password', '')
        
        if not new_password:
            return jsonify({
                'success': False,
                'error': '新密码不能为空'
            }), 400
        
        if len(new_password) < 4:
            return jsonify({
                'success': False,
                'error': '密码长度至少4位'
            }), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT 姓名 FROM users WHERE 姓名 = ?", (username,))
        if not cursor.fetchone():
            conn.close()
            return jsonify({
                'success': False,
                'error': '用户不存在'
            }), 404
        
        from werkzeug.security import generate_password_hash
        hashed_password = generate_password_hash(new_password, method='pbkdf2:sha256:10000')
        
        cursor.execute(
            "UPDATE users SET 密码 = ? WHERE 姓名 = ?",
            (hashed_password, username)
        )
        
        cursor.execute(
            "UPDATE users SET token = NULL, token_expire_time = NULL WHERE 姓名 = ?",
            (username,)
        )
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'用户 {username} 密码已更新'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'修改密码失败: {str(e)}'
        }), 500


@app.route('/api/admin/devices', methods=['GET'])
@admin_required
def get_all_devices():
    """获取所有设备绑定信息（管理员）"""
    try:
        device_manager = DeviceManager(DATABASE_PATH)
        bindings = device_manager.get_all_bindings()
        
        return jsonify({
            'success': True,
            'data': bindings
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取设备列表失败: {str(e)}'
        }), 500


@app.route('/api/admin/devices/<username>/unbind', methods=['POST'])
@admin_required
def admin_unbind_device(username):
    """解绑用户设备（管理员）"""
    try:
        device_manager = DeviceManager(DATABASE_PATH)
        success = device_manager.unbind_device(username)
        
        if success:
            return jsonify({
                'success': True,
                'message': f'用户 {username} 的设备已解绑'
            })
        else:
            return jsonify({
                'success': False,
                'error': '解绑失败'
            }), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'解绑设备失败: {str(e)}'
        }), 500


@app.route('/api/admin/login-attempts', methods=['GET'])
@admin_required
def get_failed_login_attempts():
    """获取失败的登录尝试记录（管理员）"""
    try:
        limit = request.args.get('limit', 50, type=int)
        device_manager = DeviceManager(DATABASE_PATH)
        attempts = device_manager.get_failed_attempts(limit)
        
        return jsonify({
            'success': True,
            'data': attempts
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取登录记录失败: {str(e)}'
        }), 500


@app.route('/api/leave/submit', methods=['POST'])
@login_required
def submit_leave():
    """
    提交请假申请
    
    Request Body:
    {
        "leave_type": "请假类型",
        "start_date": "开始日期",
        "end_date": "结束日期",
        "leave_days": 请假天数,
        "reason": "请假原因"
    }
    """
    try:
        data = request.get_json()
        leave_type = data.get('leave_type', '').strip()
        start_date = data.get('start_date', '').strip()
        end_date = data.get('end_date', '').strip()
        leave_days = data.get('leave_days', 0)
        reason = data.get('reason', '').strip()
        
        if not leave_type:
            return jsonify({'success': False, 'error': '请选择请假类型'}), 400
        
        if not start_date or not end_date:
            return jsonify({'success': False, 'error': '请选择请假日期'}), 400
        
        if leave_days <= 0:
            return jsonify({'success': False, 'error': '请假天数必须大于0'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            """INSERT INTO leave_records 
               (姓名, 请假类型, 开始日期, 结束日期, 请假天数, 请假原因, 申请时间)
               VALUES (?, ?, ?, ?, ?, ?, datetime('now'))""",
            (request.user_name, leave_type, start_date, end_date, leave_days, reason)
        )
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '请假申请提交成功'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'提交请假申请失败: {str(e)}'
        }), 500


@app.route('/api/leave/my-records', methods=['GET'])
@login_required
def get_my_leave_records():
    """获取当前用户的请假记录"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute(
            """SELECT * FROM leave_records 
               WHERE 姓名 = ? 
               ORDER BY 申请时间 DESC""",
            (request.user_name,)
        )
        
        records = cursor.fetchall()
        conn.close()
        
        return jsonify({
            'success': True,
            'data': [dict(record) for record in records]
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取请假记录失败: {str(e)}'
        }), 500


@app.route('/api/admin/leave/records', methods=['GET'])
@admin_required
def get_all_leave_records():
    """获取所有请假记录（管理员）"""
    try:
        name = request.args.get('name', '').strip()
        year = request.args.get('year', '').strip()
        month = request.args.get('month', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = "SELECT * FROM leave_records WHERE 1=1"
        params = []
        
        if name:
            query += " AND 姓名 LIKE ?"
            params.append(f"%{name}%")
        
        if year and month:
            query += " AND strftime('%Y', 开始日期) = ? AND strftime('%m', 开始日期) = ?"
            params.extend([year, month])
        elif year:
            query += " AND strftime('%Y', 开始日期) = ?"
            params.append(year)
        elif month:
            query += " AND strftime('%m', 开始日期) = ?"
            params.append(month)
        
        query += " ORDER BY 申请时间 DESC"
        
        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()
        
        return jsonify({
            'success': True,
            'data': [dict(record) for record in records]
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取请假记录失败: {str(e)}'
        }), 500


@app.route('/api/admin/leave/export', methods=['GET'])
@admin_required
def export_leave_records():
    """导出请假记录（管理员）"""
    try:
        import csv
        import io
        from datetime import datetime
        
        name = request.args.get('name', '').strip()
        year = request.args.get('year', '').strip()
        month = request.args.get('month', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = "SELECT * FROM leave_records WHERE 1=1"
        params = []
        
        if name:
            query += " AND 姓名 LIKE ?"
            params.append(f"%{name}%")
        
        if year and month:
            query += " AND strftime('%Y', 开始日期) = ? AND strftime('%m', 开始日期) = ?"
            params.extend([year, month])
        elif year:
            query += " AND strftime('%Y', 开始日期) = ?"
            params.append(year)
        elif month:
            query += " AND strftime('%m', 开始日期) = ?"
            params.append(month)
        
        query += " ORDER BY 申请时间 DESC"
        
        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()
        
        output = io.StringIO()
        writer = csv.writer(output, quoting=csv.QUOTE_ALL)
        
        writer.writerow([
            '姓名', '请假类型', '开始日期', '结束日期', 
            '请假天数', '请假原因', '申请时间'
        ])
        
        for record in records:
            writer.writerow([
                record['姓名'],
                record['请假类型'],
                record['开始日期'],
                record['结束日期'],
                record['请假天数'],
                record['请假原因'] or '',
                record['申请时间'] or ''
            ])
        
        output.seek(0)
        
        from flask import Response
        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename=leave_records_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            }
        )
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导出失败: {str(e)}'
        }), 500


# ========================================
# 客户档案API
# ========================================

CUSTOMER_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'customer_uploads')
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def ensure_customer_dir():
    os.makedirs(CUSTOMER_UPLOAD_DIR, exist_ok=True)


@app.route('/api/admin/customer/upload', methods=['POST'])
@admin_required
def upload_customer_excel():
    """上传客户基础表Excel文件"""
    try:
        ensure_customer_dir()
        
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': '仅支持Excel文件（.xlsx, .xls）'}), 400
        
        import pandas as pd
        
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        saved_filename = f"customer_{timestamp}_{filename}"
        filepath = os.path.join(CUSTOMER_UPLOAD_DIR, saved_filename)
        file.save(filepath)
        
        try:
            df = pd.read_excel(filepath)
        except Exception as e:
            os.remove(filepath)
            return jsonify({'success': False, 'error': f'Excel文件读取失败: {str(e)}'}), 400
        
        if df.empty:
            os.remove(filepath)
            return jsonify({'success': False, 'error': 'Excel文件内容为空'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS customer_archive (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                档案编号 TEXT,
                联系人 TEXT,
                年龄 TEXT,
                性别 TEXT,
                出生日期 TEXT,
                街道 TEXT,
                级别 TEXT,
                身份证 TEXT,
                联系电话 TEXT,
                现住址 TEXT,
                自述身高 TEXT,
                自述体重 TEXT,
                紧急联系电话 TEXT,
                居住情况 TEXT,
                慢性疾病 TEXT,
                使用药物 TEXT,
                意识 TEXT,
                生命体征 TEXT,
                四肢活动情况 TEXT,
                现在有无压疮 TEXT,
                部位 TEXT,
                压疮危险度评估 TEXT,
                日常生活活动能力评估 TEXT,
                跌倒坠床风险评估 TEXT,
                生活自理能力 TEXT,
                客户类型 TEXT,
                护理计划起始时间 TEXT,
                待遇结束时间 TEXT,
                护理计划及要点 TEXT,
                备注 TEXT,
                上传批次 TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        batch_id = timestamp
        
        columns = df.columns.tolist()
        column_mapping = {}
        for col in columns:
            col_lower = str(col).strip()
            if '档案编号' in col_lower or '编号' in col_lower:
                column_mapping[col] = '档案编号'
            elif '联系人' in col_lower or ('姓名' in col_lower and '联系' not in col_lower):
                column_mapping[col] = '联系人'
            elif '年龄' in col_lower:
                column_mapping[col] = '年龄'
            elif '性别' in col_lower:
                column_mapping[col] = '性别'
            elif '出生日期' in col_lower:
                column_mapping[col] = '出生日期'
            elif '街道' in col_lower:
                column_mapping[col] = '街道'
            elif '级别' in col_lower:
                column_mapping[col] = '级别'
            elif '身份证' in col_lower:
                column_mapping[col] = '身份证'
            elif '紧急联系电话' in col_lower or '紧急联系' in col_lower:
                column_mapping[col] = '紧急联系电话'
            elif '家庭电话' in col_lower or '老人电话' in col_lower:
                column_mapping[col] = '联系电话'
            elif '联系电话' in col_lower:
                column_mapping[col] = '联系电话'
            elif '电话' in col_lower or '手机' in col_lower:
                column_mapping[col] = '联系电话'
            elif '现住址' in col_lower or ('住址' in col_lower and '现' in col_lower):
                column_mapping[col] = '现住址'
            elif '地址' in col_lower:
                column_mapping[col] = '现住址'
            elif '身高' in col_lower:
                column_mapping[col] = '自述身高'
            elif '体重' in col_lower:
                column_mapping[col] = '自述体重'
            elif '居住情况' in col_lower or '居住' in col_lower:
                column_mapping[col] = '居住情况'
            elif '慢性疾病' in col_lower or '疾病' in col_lower:
                column_mapping[col] = '慢性疾病'
            elif '使用药物' in col_lower or '药物' in col_lower:
                column_mapping[col] = '使用药物'
            elif '意识' in col_lower:
                column_mapping[col] = '意识'
            elif '生命体征' in col_lower:
                column_mapping[col] = '生命体征'
            elif '四肢活动' in col_lower or '四肢' in col_lower:
                column_mapping[col] = '四肢活动情况'
            elif '压疮' in col_lower and '评估' not in col_lower:
                column_mapping[col] = '现在有无压疮'
            elif '部位' in col_lower or '部分' in col_lower:
                column_mapping[col] = '部位'
            elif '压疮危险度评估' in col_lower or '压疮评估' in col_lower:
                column_mapping[col] = '压疮危险度评估'
            elif '日常生活活动能力评估' in col_lower or '活动能力' in col_lower:
                column_mapping[col] = '日常生活活动能力评估'
            elif '跌倒' in col_lower or '坠床' in col_lower:
                column_mapping[col] = '跌倒坠床风险评估'
            elif '生活自理能力' in col_lower or '自理能力' in col_lower:
                column_mapping[col] = '生活自理能力'
            elif '客户类型' in col_lower or ('类型' in col_lower and '评估' not in col_lower):
                column_mapping[col] = '客户类型'
            elif '护理计划起始时间' in col_lower or ('护理计划' in col_lower and '起始' in col_lower):
                column_mapping[col] = '护理计划起始时间'
            elif '待遇结束时间' in col_lower or ('待遇' in col_lower and '结束' in col_lower):
                column_mapping[col] = '待遇结束时间'
            elif '护理计划及要点' in col_lower or ('护理计划' in col_lower and '要点' in col_lower):
                column_mapping[col] = '护理计划及要点'
            elif '备注' in col_lower or '说明' in col_lower:
                column_mapping[col] = '备注'
        
        inserted_count = 0
        for _, row in df.iterrows():
            data = {}
            for orig_col, mapped_col in column_mapping.items():
                val = row.get(orig_col)
                data[mapped_col] = str(val) if pd.notna(val) else ''
            
            remaining_cols = {k: v for k, v in row.items() if k not in column_mapping and pd.notna(v)}
            extra_info = '；'.join([f"{k}: {v}" for k, v in remaining_cols.items()]) if remaining_cols else ''
            if data.get('备注') and extra_info:
                data['备注'] = data['备注'] + '；' + extra_info
            elif extra_info:
                data['备注'] = extra_info
            
            cursor.execute("""
                INSERT INTO customer_archive (
                    档案编号, 联系人, 年龄, 性别, 出生日期, 街道, 级别, 身份证,
                    联系电话, 现住址, 自述身高, 自述体重, 紧急联系电话, 居住情况, 慢性疾病,
                    使用药物, 意识, 生命体征, 四肢活动情况, 现在有无压疮, 部位,
                    压疮危险度评估, 日常生活活动能力评估, 跌倒坠床风险评估,
                    生活自理能力, 客户类型, 护理计划起始时间, 待遇结束时间,
                    护理计划及要点, 备注, 上传批次
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                data.get('档案编号', ''),
                data.get('联系人', ''),
                data.get('年龄', ''),
                data.get('性别', ''),
                data.get('出生日期', ''),
                data.get('街道', ''),
                data.get('级别', ''),
                data.get('身份证', ''),
                data.get('联系电话', ''),
                data.get('现住址', ''),
                data.get('自述身高', ''),
                data.get('自述体重', ''),
                data.get('紧急联系电话', ''),
                data.get('居住情况', ''),
                data.get('慢性疾病', ''),
                data.get('使用药物', ''),
                data.get('意识', ''),
                data.get('生命体征', ''),
                data.get('四肢活动情况', ''),
                data.get('现在有无压疮', ''),
                data.get('部位', ''),
                data.get('压疮危险度评估', ''),
                data.get('日常生活活动能力评估', ''),
                data.get('跌倒坠床风险评估', ''),
                data.get('生活自理能力', ''),
                data.get('客户类型', ''),
                data.get('护理计划起始时间', ''),
                data.get('待遇结束时间', ''),
                data.get('护理计划及要点', ''),
                data.get('备注', ''),
                batch_id
            ))
            inserted_count += 1
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'成功导入 {inserted_count} 条客户数据',
            'batch_id': batch_id,
            'total': inserted_count
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'上传失败: {str(e)}'}), 500


@app.route('/api/admin/customer/list', methods=['GET'])
@admin_required
def get_customer_list():
    """获取客户档案列表"""
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))
        search = request.args.get('search', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = "SELECT * FROM customer_archive WHERE 1=1"
        params = []
        
        if search:
            query += " AND (档案编号 LIKE ? OR 联系人 LIKE ? OR 联系电话 LIKE ? OR 现住址 LIKE ? OR 身份证 LIKE ? OR 街道 LIKE ?)"
            search_param = f"%{search}%"
            params.extend([search_param, search_param, search_param, search_param, search_param, search_param])
        
        query += " ORDER BY created_at DESC"
        
        cursor.execute(query, params)
        all_records = cursor.fetchall()
        
        total = len(all_records)
        start = (page - 1) * per_page
        end = start + per_page
        records = all_records[start:end]
        
        result = []
        for record in records:
            result.append({
                'id': record['id'],
                '档案编号': record['档案编号'] or '',
                '联系人': record['联系人'] or '',
                '年龄': record['年龄'] or '',
                '性别': record['性别'] or '',
                '出生日期': record['出生日期'] or '',
                '街道': record['街道'] or '',
                '级别': record['级别'] or '',
                '身份证': record['身份证'] or '',
                '联系电话': record['联系电话'] or '',
                '现住址': record['现住址'] or '',
                '自述身高': record['自述身高'] or '',
                '自述体重': record['自述体重'] or '',
                '居住情况': record['居住情况'] or '',
                '慢性疾病': record['慢性疾病'] or '',
                '使用药物': record['使用药物'] or '',
                '意识': record['意识'] or '',
                '生命体征': record['生命体征'] or '',
                '四肢活动情况': record['四肢活动情况'] or '',
                '现在有无压疮': record['现在有无压疮'] or '',
                '压疮危险度评估': record['压疮危险度评估'] or '',
                '日常生活活动能力评估': record['日常生活活动能力评估'] or '',
                '跌倒坠床风险评估': record['跌倒坠床风险评估'] or '',
                '生活自理能力': record['生活自理能力'] or '',
                '客户类型': record['客户类型'] or '',
                '护理计划起始时间': record['护理计划起始时间'] or '',
                '待遇结束时间': record['待遇结束时间'] or '',
                '护理计划及要点': record['护理计划及要点'] or '',
                '备注': record['备注'] or '',
                '上传批次': record['上传批次'] or '',
                'created_at': record['created_at'] or ''
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'data': result,
            'total': total,
            'page': page,
            'per_page': per_page
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'获取客户列表失败: {str(e)}'}), 500


@app.route('/api/admin/customer/delete/<int:customer_id>', methods=['DELETE'])
@admin_required
def delete_customer(customer_id):
    """删除客户档案记录"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM customer_archive WHERE id = ?", (customer_id,))
        conn.commit()
        affected = cursor.rowcount
        conn.close()
        
        if affected > 0:
            return jsonify({'success': True, 'message': '删除成功'})
        else:
            return jsonify({'success': False, 'error': '记录不存在'}), 404
            
    except Exception as e:
        return jsonify({'success': False, 'error': f'删除失败: {str(e)}'}), 500


@app.route('/api/admin/customer/delete-all', methods=['DELETE'])
@admin_required
def delete_all_customers():
    """删除所有客户档案记录"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM customer_archive")
        count = cursor.fetchone()[0]
        
        if count == 0:
            conn.close()
            return jsonify({'success': False, 'error': '没有客户数据可删除'}), 400
        
        cursor.execute("DELETE FROM customer_archive")
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': f'已删除 {count} 条客户记录'})
            
    except Exception as e:
        return jsonify({'success': False, 'error': f'删除失败: {str(e)}'}), 500


@app.route('/api/admin/customer/export/start', methods=['POST'])
@admin_required
def start_export_task():
    """启动异步导出任务"""
    try:
        search = request.json.get('search', '').strip() if request.json else ''
        
        task_id = str(uuid.uuid4())
        
        _create_export_task(task_id, {
            'status': 'pending',
            'progress': 0,
            'total': 0,
            'message': '正在准备导出...',
            'file_path': None,
            'filename': None,
            'error': None,
            'created_at': datetime.now().isoformat()
        })
        
        thread = threading.Thread(target=do_export_task, args=(task_id, search))
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'task_id': task_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/admin/customer/export/progress/<task_id>', methods=['GET'])
@admin_required
def get_export_progress(task_id):
    """获取导出任务进度"""
    task = _get_export_task(task_id)
    if not task:
        return jsonify({'success': False, 'error': '任务不存在'}), 404
    
    return jsonify({
        'success': True,
        'status': task['status'],
        'progress': task['progress'],
        'total': task['total'],
        'message': task['message'],
        'file_path': task['file_path'],
        'filename': task['filename'],
        'error': task['error']
    })


@app.route('/api/admin/customer/export/download/<task_id>', methods=['GET'])
@admin_required
def download_export_file(task_id):
    """下载导出的文件"""
    task = _get_export_task(task_id)
    if not task:
        return jsonify({'success': False, 'error': '任务不存在'}), 404
    
    if task['status'] != 'completed':
        return jsonify({'success': False, 'error': '任务未完成'}), 400
    
    if not task['file_path'] or not os.path.exists(task['file_path']):
        return jsonify({'success': False, 'error': '文件不存在'}), 404
    
    from urllib.parse import quote
    filename = task['filename']
    encoded_filename = quote(filename)
    
    mimetype = 'application/zip' if filename.endswith('.zip') else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    response = send_file(
        task['file_path'],
        as_attachment=True,
        download_name=filename,
        mimetype=mimetype
    )
    response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
    return response


def do_export_task(task_id, search):
    """执行导出任务的后台函数"""
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import zipfile
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        _set_export_task(task_id, status='processing', message='正在查询数据...')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = "SELECT * FROM customer_archive WHERE 1=1"
        params = []
        
        if search:
            query += " AND (档案编号 LIKE ? OR 联系人 LIKE ? OR 联系电话 LIKE ? OR 现住址 LIKE ? OR 身份证 LIKE ? OR 街道 LIKE ?)"
            search_param = f"%{search}%"
            params.extend([search_param, search_param, search_param, search_param, search_param, search_param])
        
        query += " ORDER BY created_at DESC"
        
        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()
        
        if not records:
            _set_export_task(task_id, status='error', error='没有客户数据可导出')
            return
        
        _set_export_task(task_id, total=len(records), message=f'正在生成 {len(records)} 份文档...')
        
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'logo.png')
        if not os.path.exists(logo_path):
            print(f"logo图片不存在: {logo_path}")
            logo_path = None
        else:
            print(f"logo图片路径: {logo_path}")
        
        output_dir = os.path.join(CUSTOMER_UPLOAD_DIR, 'exports')
        os.makedirs(output_dir, exist_ok=True)
        
        temp_dir = os.path.join(output_dir, f'temp_{task_id}')
        if os.path.exists(temp_dir):
            for f in os.listdir(temp_dir):
                try:
                    os.remove(os.path.join(temp_dir, f))
                except:
                    pass
        os.makedirs(temp_dir, exist_ok=True)
        
        def get_field(record, field_name):
            try:
                val = record[field_name]
                if val is None:
                    return ''
                result = str(val) if val else ''
                if result.endswith('.0'):
                    result = result[:-2]
                return result
            except (KeyError, TypeError, IndexError):
                return ''
        
        def format_date_only(date_str):
            if not date_str:
                return ''
            try:
                if ' ' in date_str:
                    return date_str.split(' ')[0]
                if 'T' in date_str:
                    return date_str.split('T')[0]
                return date_str
            except:
                return date_str
        
        def set_run_font(run, font_name='黑体', font_size=11, bold=True, underline=False):
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.underline = underline
            try:
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font_name)
            except Exception as e:
                pass
        
        import threading
        progress_lock = threading.Lock()
        completed_count = [0]
        
        def generate_single_doc(args):
            idx, record, temp_dir, logo_path, task_id = args
            try:
                from docx import Document
                from docx.shared import Cm, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                
                doc = Document()
                
                section = doc.sections[0]
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(4.0)
                
                para0 = doc.add_paragraph()
                para0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para0.paragraph_format.space_before = Pt(4.25)
                para0.paragraph_format.line_spacing = 0.8
                run0 = para0.add_run("编号：")
                run0.font.name = '黑体'
                run0.font.size = Pt(14)
                run0.font.bold = True
                try:
                    r = run0._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
                run0_val = para0.add_run(get_field(record, '档案编号'))
                run0_val.font.name = '黑体'
                run0_val.font.size = Pt(14)
                run0_val.font.bold = True
                run0_val.font.underline = True
                try:
                    r = run0_val._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
                
                if logo_path:
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        img_para.paragraph_format.space_before = Pt(0)
                        img_run = img_para.add_run()
                        img_run.add_picture(logo_path, width=Cm(4.37), height=Cm(1.28))
                    except Exception as e:
                        print(f"异步导出-插入图片失败: {e}")
                
                para2 = doc.add_paragraph()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para2.paragraph_format.space_before = Pt(6.25)
                para2.paragraph_format.line_spacing = 0.8
                run2 = para2.add_run('长护险客户基本信息及护理计划')
                run2.font.name = '黑体'
                run2.font.size = Pt(16)
                run2.font.bold = True
                try:
                    r = run2._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
                
                para3 = doc.add_paragraph()
                para3.paragraph_format.line_spacing = 1.62
                
                def add_content_para(parts, space_before=3.9, left_indent=0, first_line_indent=0):
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(space_before)
                    para.paragraph_format.line_spacing = 0.97
                    if left_indent > 0:
                        para.paragraph_format.left_indent = Cm(left_indent)
                    if first_line_indent > 0:
                        para.paragraph_format.first_line_indent = Cm(first_line_indent)
                    for text, is_underline in parts:
                        run = para.add_run(text)
                        run.font.name = '黑体'
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.underline = is_underline
                        try:
                            r = run._element
                            rPr = r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), '黑体')
                        except:
                            pass
                    return para
                
                add_content_para([
                    ("姓名：", False), (get_field(record, '联系人'), True),
                    ("           年龄：", False), (get_field(record, '年龄'), True),
                    ("                 岁      性别：", False), (get_field(record, '性别'), True)
                ], space_before=3.9)
                
                add_content_para([
                    ("出生日期：", False), (format_date_only(get_field(record, '出生日期')), True),
                    ("             街道：", False), (get_field(record, '街道'), True),
                    ("                    级别：", False), (get_field(record, '级别'), True)
                ], space_before=10.35, left_indent=0.05)
                
                add_content_para([
                    ("身份证号：", False), (get_field(record, '身份证'), True),
                    ("          家庭电话(老人电话):", False), (get_field(record, '联系电话'), True)
                ], space_before=9.4, left_indent=0.05)
                
                add_content_para([
                    ("现住址：", False), (get_field(record, '现住址'), True),
                    ("                   医保类型:", False), (get_field(record, '客户类型'), True)
                ], space_before=9.4, left_indent=0.05)
                
                add_content_para([
                    ("自述身高：", False), (get_field(record, '自述身高'), True),
                    ("CM    自述体重：", False), (get_field(record, '自述体重'), True),
                    ("KG   紧急联系电话：", False), (get_field(record, '紧急联系电话'), True)
                ], space_before=3.75)
                
                add_content_para([("居住情况：", False), (get_field(record, '居住情况'), True)], space_before=8.8, left_indent=0.06)
                add_content_para([("慢性疾病：", False), (get_field(record, '慢性疾病'), True)], space_before=8.8, left_indent=0.06)
                add_content_para([("使用药物：", False), (get_field(record, '使用药物'), True)], space_before=8.8)
                
                add_content_para([("意识：", False), (get_field(record, '意识'), True)], space_before=8.8)
                add_content_para([("生命体征：", False), (get_field(record, '生命体征'), True)], space_before=8.8)
                
                add_content_para([("四肢活动情况：", False), (get_field(record, '四肢活动情况'), True)], space_before=11.2)
                
                add_content_para([
                    ("现在有无压疮：", False), (get_field(record, '现在有无压疮'), True),
                    ("                 部位：", False), (get_field(record, '部位'), True)
                ], space_before=12.3, left_indent=0.2)
                
                add_content_para([
                    ("压疮危险度评估：", False), (get_field(record, '压疮危险度评估'), True),
                    ("    分", False),
                    ("（(Braden评分<18分，提示有发生压疮的风险）", False)
                ], space_before=14.6, first_line_indent=2.45)
                
                add_content_para([
                    ("日常生活活动能力评估（ADI总分）：", False), (get_field(record, '日常生活活动能力评估'), True),
                    ("     分", False)
                ], space_before=3.0, first_line_indent=2.45)
                
                add_content_para([
                    ("跌倒/坠床风险评估(如果正常分数为0)：", False), (get_field(record, '跌倒坠床风险评估'), True),
                    ("   分", False),
                    ("（3分以上有跌倒风险）", False)
                ], space_before=3.0, first_line_indent=2.45)
                
                add_content_para([("生活自理能力：", False), (get_field(record, '生活自理能力'), True)], space_before=13.7)
                
                add_content_para([
                    ("护理计划起始时间：", False), (format_date_only(get_field(record, '护理计划起始时间')), True),
                    ("                待遇结束时间：", False), (format_date_only(get_field(record, '待遇结束时间')), True)
                ], space_before=14.0, left_indent=0.19)
                
                add_content_para([("护理计划及要点：", False), (get_field(record, '护理计划及要点'), True)], space_before=17.1, left_indent=0.19)
                
                para_end1 = doc.add_paragraph()
                para_end1.paragraph_format.space_before = Pt(20)
                
                para_end2 = doc.add_paragraph()
                para_end2.paragraph_format.space_before = Pt(15)
                para_end2.paragraph_format.left_indent = Cm(0.64)
                para_end2.paragraph_format.line_spacing = 0.96
                run_end = para_end2.add_run("护士/评估者签字：                      客户或家属签字：")
                run_end.font.name = '黑体'
                run_end.font.size = Pt(11)
                run_end.font.bold = True
                try:
                    r = run_end._element
                    rPr = r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
                
                name = get_field(record, '联系人') or '未知'
                number = get_field(record, '档案编号')
                record_id = get_field(record, 'id') or str(idx + 1)
                filename = f"{name}+{number}.docx" if number else f"{name}_{record_id}.docx"
                filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
                
                output_path = os.path.join(temp_dir, filename)
                doc.save(output_path)
                del doc
                
                with progress_lock:
                    completed_count[0] += 1
                    _set_export_task(task_id, progress=completed_count[0], message=f'正在生成文档 ({completed_count[0]}/{len(records)})...')
                
                return (output_path, filename)
                    
            except Exception as record_e:
                print(f"处理记录 {idx + 1} 失败: {record_e}")
                with progress_lock:
                    completed_count[0] += 1
                    _set_export_task(task_id, progress=completed_count[0])
                return None
        
        tasks = [(idx, record, temp_dir, logo_path, task_id) for idx, record in enumerate(records)]
        
        doc_files = []
        max_workers = min(os.cpu_count() or 4, len(records))
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(generate_single_doc, task) for task in tasks]
            for future in as_completed(futures):
                result = future.result()
                if result:
                    doc_files.append(result)
        
        if not doc_files:
            _set_export_task(task_id, status='error', error='没有成功生成任何文档')
            return
        
        _set_export_task(task_id, message='正在打包文件...')
        
        if len(doc_files) == 1:
            output_path, output_filename = doc_files[0]
            _set_export_task(task_id, status='completed', file_path=output_path, filename=output_filename, message='导出完成，可以下载')
        else:
            zip_filename = "客户档案表.zip"
            zip_path = os.path.join(output_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_STORED) as zipf:
                for file_path, filename in doc_files:
                    if os.path.exists(file_path):
                        zipf.write(file_path, filename)
            
            for file_path, _ in doc_files:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except:
                    pass
            try:
                os.rmdir(temp_dir)
            except:
                pass
            
            _set_export_task(task_id, status='completed', file_path=zip_path, filename=zip_filename, message='导出完成，可以下载')
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        _set_export_task(task_id, status='error', error=f'导出失败: {str(e)}')


@app.route('/api/admin/customer/export-word', methods=['GET'])
@admin_required
def export_customer_word():
    """导出客户档案纸质表为Word文档"""
    try:
        from docx import Document
        from docx.shared import Cm, Pt, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import zipfile
        import shutil
        import gc
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from io import BytesIO
        
        search = request.args.get('search', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = "SELECT * FROM customer_archive WHERE 1=1"
        params = []
        
        if search:
            query += " AND (档案编号 LIKE ? OR 联系人 LIKE ? OR 联系电话 LIKE ? OR 现住址 LIKE ? OR 身份证 LIKE ? OR 街道 LIKE ?)"
            search_param = f"%{search}%"
            params.extend([search_param, search_param, search_param, search_param, search_param, search_param])
        
        query += " ORDER BY created_at DESC"
        
        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()
        
        if not records:
            return jsonify({'success': False, 'error': '没有客户数据可导出'}), 400
        
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'logo.png')
        if not os.path.exists(logo_path):
            print(f"同步导出-logo图片不存在: {logo_path}")
            logo_path = None
        else:
            print(f"同步导出-logo图片路径: {logo_path}")
        
        output_dir = os.path.join(CUSTOMER_UPLOAD_DIR, 'exports')
        os.makedirs(output_dir, exist_ok=True)
        
        temp_dir = os.path.join(output_dir, 'temp_docs')
        if os.path.exists(temp_dir):
            for f in os.listdir(temp_dir):
                try:
                    os.remove(os.path.join(temp_dir, f))
                except:
                    pass
        os.makedirs(temp_dir, exist_ok=True)
        
        def get_field(record, field_name):
            try:
                val = record[field_name]
                if val is None:
                    return ''
                result = str(val) if val else ''
                if result.endswith('.0'):
                    result = result[:-2]
                return result
            except (KeyError, TypeError, IndexError):
                return ''
        
        def format_date_only(date_str):
            if not date_str:
                return ''
            try:
                if ' ' in date_str:
                    return date_str.split(' ')[0]
                if 'T' in date_str:
                    return date_str.split('T')[0]
                return date_str
            except:
                return date_str
        
        def set_run_font(run, font_name='黑体', font_size=11, bold=True, underline=False):
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.underline = underline
            try:
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font_name)
            except Exception as e:
                pass
        
        def generate_single_doc(args):
            idx, record, temp_dir, logo_image_data = args
            try:
                doc = Document()
                
                section = doc.sections[0]
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(4.0)
                
                para0 = doc.add_paragraph()
                para0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para0.paragraph_format.space_before = Pt(4.25)
                para0.paragraph_format.line_spacing = 0.8
                run0 = para0.add_run("编号：")
                set_run_font(run0, font_size=14, bold=True)
                run0_val = para0.add_run(get_field(record, '档案编号'))
                set_run_font(run0_val, font_size=14, bold=True, underline=True)
                
                if logo_path:
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        img_para.paragraph_format.space_before = Pt(0)
                        img_run = img_para.add_run()
                        img_run.add_picture(logo_path, width=Cm(4.37), height=Cm(1.28))
                    except Exception as e:
                        print(f"同步导出-插入图片失败: {e}")
                        import traceback
                        traceback.print_exc()
                
                para2 = doc.add_paragraph()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para2.paragraph_format.space_before = Pt(6.25)
                para2.paragraph_format.line_spacing = 0.8
                run2 = para2.add_run('长护险客户基本信息及护理计划')
                set_run_font(run2, font_size=16, bold=True)
                
                para3 = doc.add_paragraph()
                para3.paragraph_format.line_spacing = 1.62
                
                def add_content_para(parts, space_before=3.9, left_indent=0, first_line_indent=0):
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(space_before)
                    para.paragraph_format.line_spacing = 0.97
                    if left_indent > 0:
                        para.paragraph_format.left_indent = Cm(left_indent)
                    if first_line_indent > 0:
                        para.paragraph_format.first_line_indent = Cm(first_line_indent)
                    for text, is_underline in parts:
                        run = para.add_run(text)
                        set_run_font(run, underline=is_underline)
                    return para
                
                add_content_para([
                    ("姓名：", False), (get_field(record, '联系人'), True),
                    ("           年龄：", False), (get_field(record, '年龄'), True),
                    ("                 岁      性别：", False), (get_field(record, '性别'), True)
                ], space_before=3.9)
                
                add_content_para([
                    ("出生日期：", False), (format_date_only(get_field(record, '出生日期')), True),
                    ("             街道：", False), (get_field(record, '街道'), True),
                    ("                    级别：", False), (get_field(record, '级别'), True)
                ], space_before=10.35, left_indent=0.05)
                
                add_content_para([
                    ("身份证号：", False), (get_field(record, '身份证'), True),
                    ("          家庭电话(老人电话):", False), (get_field(record, '联系电话'), True)
                ], space_before=9.4, left_indent=0.05)
                
                add_content_para([
                    ("现住址：", False), (get_field(record, '现住址'), True),
                    ("                   医保类型:", False), (get_field(record, '客户类型'), True)
                ], space_before=9.4, left_indent=0.05)
                
                add_content_para([
                    ("自述身高：", False), (get_field(record, '自述身高'), True),
                    ("CM    自述体重：", False), (get_field(record, '自述体重'), True),
                    ("KG   紧急联系电话：", False), (get_field(record, '紧急联系电话'), True)
                ], space_before=3.75)
                
                add_content_para([("居住情况：", False), (get_field(record, '居住情况'), True)], space_before=8.8, left_indent=0.06)
                add_content_para([("慢性疾病：", False), (get_field(record, '慢性疾病'), True)], space_before=8.8, left_indent=0.06)
                add_content_para([("使用药物：", False), (get_field(record, '使用药物'), True)], space_before=8.8)
                
                add_content_para([("意识：", False), (get_field(record, '意识'), True)], space_before=8.8)
                add_content_para([("生命体征：", False), (get_field(record, '生命体征'), True)], space_before=8.8)
                
                add_content_para([("四肢活动情况：", False), (get_field(record, '四肢活动情况'), True)], space_before=11.2)
                
                add_content_para([
                    ("现在有无压疮：", False), (get_field(record, '现在有无压疮'), True),
                    ("                 部位：", False), (get_field(record, '部位'), True)
                ], space_before=12.3, left_indent=0.2)
                
                add_content_para([
                    ("压疮危险度评估：", False), (get_field(record, '压疮危险度评估'), True),
                    ("    分", False),
                    ("（(Braden评分<18分，提示有发生压疮的风险）", False)
                ], space_before=14.6, first_line_indent=2.45)
                
                add_content_para([
                    ("日常生活活动能力评估（ADI总分）：", False), (get_field(record, '日常生活活动能力评估'), True),
                    ("     分", False)
                ], space_before=3.0, first_line_indent=2.45)
                
                add_content_para([
                    ("跌倒/坠床风险评估(如果正常分数为0)：", False), (get_field(record, '跌倒坠床风险评估'), True),
                    ("   分", False),
                    ("（3分以上有跌倒风险）", False)
                ], space_before=3.0, first_line_indent=2.45)
                
                add_content_para([("生活自理能力：", False), (get_field(record, '生活自理能力'), True)], space_before=13.7)
                
                add_content_para([
                    ("护理计划起始时间：", False), (format_date_only(get_field(record, '护理计划起始时间')), True),
                    ("                待遇结束时间：", False), (format_date_only(get_field(record, '待遇结束时间')), True)
                ], space_before=14.0, left_indent=0.19)
                
                add_content_para([("护理计划及要点：", False), (get_field(record, '护理计划及要点'), True)], space_before=17.1, left_indent=0.19)
                
                para_end1 = doc.add_paragraph()
                para_end1.paragraph_format.space_before = Pt(20)
                
                para_end2 = doc.add_paragraph()
                para_end2.paragraph_format.space_before = Pt(15)
                para_end2.paragraph_format.left_indent = Cm(0.64)
                para_end2.paragraph_format.line_spacing = 0.96
                run_end = para_end2.add_run("护士/评估者签字：                      客户或家属签字：")
                set_run_font(run_end)
                
                name = get_field(record, '联系人') or '未知'
                number = get_field(record, '档案编号')
                record_id = get_field(record, 'id') or str(idx + 1)
                filename = f"{name}+{number}.docx" if number else f"{name}_{record_id}.docx"
                filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
                
                output_path = os.path.join(temp_dir, filename)
                doc.save(output_path)
                del doc
                return (output_path, filename)
                    
            except Exception as record_e:
                print(f"处理记录 {idx + 1} 失败: {record_e}")
                return None
        
        tasks = [(idx, record, temp_dir, logo_image_data) for idx, record in enumerate(records)]
        
        doc_files = []
        max_workers = min(16, len(records))
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(generate_single_doc, task) for task in tasks]
            for future in as_completed(futures):
                result = future.result()
                if result:
                    doc_files.append(result)
        
        if not doc_files:
            return jsonify({'success': False, 'error': '没有成功生成任何文档'}), 400
        
        if len(doc_files) == 1:
            output_path, output_filename = doc_files[0]
            from urllib.parse import quote
            encoded_filename = quote(output_filename)
            response = send_file(
                output_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
            return response
        else:
            zip_filename = "客户档案表.zip"
            zip_path = os.path.join(output_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_STORED) as zipf:
                for file_path, filename in doc_files:
                    if os.path.exists(file_path):
                        zipf.write(file_path, filename)
            
            for file_path, _ in doc_files:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except:
                    pass
            try:
                os.rmdir(temp_dir)
            except:
                pass
            
            response = send_file(
                zip_path,
                as_attachment=True,
                download_name=zip_filename,
                mimetype='application/zip'
            )
            response.headers['Content-Disposition'] = f"attachment; filename={zip_filename}"
            return response
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'导出失败: {str(e)}'}), 500


@app.route('/api/admin/cleanup', methods=['POST'])
@admin_required
def cleanup_data():
    """清理数据"""
    try:
        import shutil
        
        data = request.get_json()
        cleanup_type = data.get('type', 'all')
        
        messages = []
        
        if cleanup_type in ['salary', 'all']:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM salary_table")
            deleted_rows = cursor.rowcount
            conn.commit()
            conn.close()
            messages.append(f'工资表：已清空 {deleted_rows} 条记录')
        
        if cleanup_type in ['customer', 'all']:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM customer_archive")
            deleted_rows = cursor.rowcount
            conn.commit()
            conn.close()
            messages.append(f'客户档案表：已清空 {deleted_rows} 条记录')
        
        if cleanup_type in ['signature', 'all']:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM signature_files")
            deleted_rows = cursor.rowcount
            cursor.execute("DELETE FROM summary_table")
            summary_deleted = cursor.rowcount
            conn.commit()
            conn.close()
            messages.append(f'签收记录：已清空 {deleted_rows} 条签名文件记录')
            messages.append(f'签收汇总：已清空 {summary_deleted} 条签收汇总记录')
            
            if os.path.exists(SIGNATURE_DIR):
                deleted_files = 0
                for filename in os.listdir(SIGNATURE_DIR):
                    file_path = os.path.join(SIGNATURE_DIR, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                            deleted_files += 1
                        elif os.path.isdir(file_path):
                            import shutil
                            shutil.rmtree(file_path)
                    except Exception as e:
                        print(f"[WARN] 删除文件失败: {file_path}, {e}")
                messages.append(f'签名文件目录已清空（删除 {deleted_files} 个文件）')
        
        if cleanup_type in ['leave', 'all']:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM leave_records")
            deleted_rows = cursor.rowcount
            conn.commit()
            conn.close()
            messages.append(f'请假记录：已清空 {deleted_rows} 条记录')
        
        if cleanup_type in ['cache', 'all']:
            temp_dir = os.path.join(os.path.dirname(__file__), 'temp')
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                os.makedirs(temp_dir)
                messages.append('临时文件目录已清空')
            
            global export_tasks
            export_tasks = {}
            messages.append('导出任务记录已清空')
        
        return jsonify({
            'success': True,
            'message': '清理完成\n' + '\n'.join(messages)
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'清理失败: {str(e)}'}), 500


# ========================================
# APP版本管理API
# ========================================

APK_DIR = os.path.join(os.path.dirname(__file__), 'apk_files')
os.makedirs(APK_DIR, exist_ok=True)


def init_app_version_table():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS app_versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            version TEXT NOT NULL,
            version_name TEXT,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_size INTEGER,
            update_note TEXT,
            upload_time DATETIME DEFAULT CURRENT_TIMESTAMP,
            is_active INTEGER DEFAULT 1
        )
    ''')
    conn.commit()
    conn.close()


init_app_version_table()


@app.route('/api/app-version/upload', methods=['POST'])
@admin_required
def upload_app_version():
    try:
        if 'apk' not in request.files:
            return jsonify({'success': False, 'error': '未找到APK文件'}), 400
        
        file = request.files['apk']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400
        
        if not file.filename.endswith('.apk'):
            return jsonify({'success': False, 'error': '只支持APK文件'}), 400
        
        version = request.form.get('version', '').strip()
        if not version:
            return jsonify({'success': False, 'error': '请输入版本号'}), 400
        
        update_note = request.form.get('update_note', '').strip()
        
        file_name = f'app_v{version}.apk'
        file_path = os.path.join(APK_DIR, file_name)
        
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("UPDATE app_versions SET is_active = 0")
        
        cursor.execute('''
            INSERT INTO app_versions (version, version_name, file_name, file_path, file_size, update_note, is_active)
            VALUES (?, ?, ?, ?, ?, ?, 1)
        ''', (version, version, file_name, file_path, file_size, update_note))
        
        cursor.execute('''
            DELETE FROM app_versions 
            WHERE is_active = 0 
            AND id NOT IN (SELECT id FROM app_versions ORDER BY upload_time DESC LIMIT 1)
        ''')
        
        for old_file in cursor.execute("SELECT file_path FROM app_versions WHERE is_active = 0").fetchall():
            old_path = old_file[0]
            if os.path.exists(old_path) and old_path != file_path:
                try:
                    os.remove(old_path)
                except:
                    pass
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'APK上传成功，版本 {version}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'上传失败: {str(e)}'}), 500


@app.route('/api/app-version/latest', methods=['GET'])
def get_latest_app_version():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT version, version_name, file_name, file_size, update_note, upload_time
            FROM app_versions 
            WHERE is_active = 1 
            ORDER BY upload_time DESC 
            LIMIT 1
        ''')
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            return jsonify({
                'success': True,
                'data': {
                    'version': row[0],
                    'version_name': row[1],
                    'file_name': row[2],
                    'file_size': row[3],
                    'update_note': row[4],
                    'upload_time': row[5]
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': '暂无APK文件'
            })
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/app-version/check', methods=['GET'])
def check_app_update():
    try:
        current_version = request.args.get('version', '0.0.0')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT version, version_name, update_note, upload_time
            FROM app_versions 
            WHERE is_active = 1 
            ORDER BY upload_time DESC 
            LIMIT 1
        ''')
        
        row = cursor.fetchone()
        conn.close()
        
        if row:
            latest_version = row[0]
            
            def parse_version(v):
                parts = v.split('.')
                return tuple(int(p) for p in parts if p.isdigit())
            
            try:
                current = parse_version(current_version)
                latest = parse_version(latest_version)
                
                if latest > current:
                    return jsonify({
                        'success': True,
                        'need_update': True,
                        'latest_version': latest_version,
                        'current_version': current_version,
                        'update_note': row[2],
                        'download_url': '/api/app-version/download'
                    })
                else:
                    return jsonify({
                        'success': True,
                        'need_update': False,
                        'latest_version': latest_version,
                        'current_version': current_version
                    })
            except:
                return jsonify({
                    'success': True,
                    'need_update': False,
                    'latest_version': latest_version,
                    'current_version': current_version
                })
        else:
            return jsonify({
                'success': True,
                'need_update': False,
                'message': '暂无可用版本'
            })
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/app-version/download', methods=['GET'])
def download_app_version():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT file_path, file_name 
            FROM app_versions 
            WHERE is_active = 1 
            ORDER BY upload_time DESC 
            LIMIT 1
        ''')
        
        row = cursor.fetchone()
        conn.close()
        
        if row and os.path.exists(row[0]):
            return send_file(row[0], as_attachment=True, download_name=row[1])
        else:
            return jsonify({'success': False, 'error': 'APK文件不存在'}), 404
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/app-version/delete', methods=['POST'])
@admin_required
def delete_app_version():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT file_path FROM app_versions WHERE is_active = 1
        ''')
        
        rows = cursor.fetchall()
        
        for row in rows:
            if os.path.exists(row[0]):
                os.remove(row[0])
        
        cursor.execute("DELETE FROM app_versions")
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': 'APK已删除'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    print("=" * 60)
    print("[INFO] 人事管理系统后端服务启动")
    print("=" * 60)
    print(f"数据库路径: {DATABASE_PATH}")
    print(f"签名目录: {SIGNATURE_DIR}")
    print("\nWeb界面:")
    print("  - GET  / - 首页（重定向到登录页）")
    print("  - GET  /login - 管理员登录页面")
    print("  - GET  /admin - 管理员后台页面")
    print("\nAPI端点:")
    print("  用户API:")
    print("    - POST /api/login - 用户登录")
    print("    - POST /api/logout - 用户登出")
    print("    - POST /api/change-password - 修改密码")
    print("    - GET  /api/verify-token - 验证Token")
    print("    - GET  /api/salary - 获取工资信息（自动更新查询状态）")
    print("    - POST /api/signature - 上传签名（支持JPG/PNG/PDF，≤10MB）")
    print("    - GET  /api/signature/<file_id> - 获取签名文件")
    print("    - GET  /api/signatures - 获取当前用户的所有签名文件")
    print("\n  管理员API:")
    print("    - GET  /api/admin/users - 获取所有用户状态")
    print("    - POST /api/admin/resign - 要求重新签名")
    print("    - GET  /api/admin/statistics - 获取查询统计信息")
    print("    - GET  /api/admin/department-statistics - 获取部门统计信息")
    print("    - GET  /api/admin/summaries - 获取所有汇总数据")
    print("    - GET  /api/admin/signatures - 获取所有签名文件")
    print("    - GET  /api/admin/signature-statistics - 获取签名文件统计信息")
    print("    - POST /api/admin/sync-all - 同步所有工资数据到汇总表")
    print("    - POST /api/admin/import-csv - 从CSV导入工资数据")
    print("    - POST /api/admin/upload-salary - 上传工资表文件（支持Excel和CSV）")
    print("    - GET  /api/admin/export-csv - 导出汇总数据到CSV")
    print("    - GET  /api/admin/export-report - 导出统计报表")
    print("    - GET  /api/admin/data-summary - 获取数据汇总信息")
    print("\n默认账号:")
    print("  - admin / admin123 (管理员)")
    print("  - 张三 / 123456")
    print("  - 李四 / 123456")
    print("  - 王五 / 123456")
    print("=" * 60)
    print("[INFO] 服务支持IPv4和IPv6访问")
    print("[INFO] Web管理界面: http://localhost:32996/login")
    print("=" * 60)
    
    app.run(host='::', port=32996, debug=False)
